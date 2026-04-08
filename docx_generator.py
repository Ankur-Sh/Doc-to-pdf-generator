import os
import subprocess
import time
import warnings
import platform
import docx2pdf
import base64
from io import BytesIO
from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from parsing_state import ParsingState
from text_helper import bold
from pdf_transformer import convert_all_pdfs
from docx_reader import read_file_content

# Suppress warnings
warnings.filterwarnings('ignore', category=UserWarning, module='docx.styles')

# Windows-specific utilities
def is_windows():
    return platform.system() == "Windows"

def close_word_if_needed():
    """Automatically close Word if on Windows"""
    if is_windows():
        try:
            subprocess.run(["taskkill", "/F", "/IM", "WINWORD.EXE"], 
                         capture_output=True, timeout=3, shell=True)
            time.sleep(0.5)
        except:
            pass

def safe_file_operation(file_path, operation, max_retries=3):
    """Safely perform file operations with automatic retry"""
    file_path = os.path.normpath(file_path)
    for attempt in range(max_retries):
        try:
            if attempt > 0 and is_windows():
                close_word_if_needed()
                time.sleep(0.5)
            return operation(file_path)
        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(1)
                continue
            raise
    return None


font_name = "Source Sans Pro"
font_size = 12
line_spacing = 1.5
BOLD_ITALIC_IDENTIFIER = {
    "b": "b",
    "i": "i"
}

def get_text_width(document):
    section = document.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 36000

def set_cell_background(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def __check_state(question):
    if (len(question["question"]) == 0):
        return (-1, "Error Type: Question")
    
    if (len(question["options"]) == 0):
        return (-1, "Error Type: Options")

    
    if (question["answer"] == -1):
        return (-1, "Error Type: Answer")

    
    if (len(question["explanation"]) == 0):
        return (-1, "Error Type: Explanation")

    
    return (1, None)

def check_state(questions: list):
    idx = 0 
    errors = False
    hasErrors = False
    for question in questions:
        idx += 1
        errors, error_type = __check_state(question)
        if (errors < 0):
            # Only show error if it's critical (missing question/options)
            # For missing answers/explanations, just warn but continue
            if error_type in ["Error Type: Question", "Error Type: Options"]:
                print(f"There is a error with this Question, {error_type=}\n{question}")
                hasErrors = True
            elif error_type == "Error Type: Answer":
                # Try to auto-detect answer from options if possible
                # Skip this question if we can't find answer
                print(f"⚠ Warning: Question {idx} missing answer - skipping")
                continue
            elif error_type == "Error Type: Explanation":
                # Allow questions without explanation (set empty)
                question["explanation"] = "No explanation provided."

    if (hasErrors):
        print("Critical errors found. Please fix and try again.")
        exit()        

def add_empty_paragraph(document):
    return document.add_paragraph()

def add_image(document: Document, image_data: str):
    """Add image to document from base64 data"""
    try:
        if not image_data:
            print("⚠ Warning: Empty image data provided")
            return
            
        if 'base64,' in image_data:
            base64_str = image_data.split(',', 1)[1]
        else:
            base64_str = image_data
        
        # Remove any trailing characters
        base64_str = base64_str.strip()
        
        if not base64_str:
            print("⚠ Warning: Empty base64 string after processing")
            return

        # Decode base64
        try:
            image_bytes = base64.b64decode(base64_str)
            if len(image_bytes) == 0:
                print("⚠ Warning: Decoded image is empty")
                return
        except Exception as e:
            print(f"⚠ Warning: Could not decode base64: {e}")
            return

        image_stream = BytesIO(image_bytes)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        
        # Calculate width
        width = Mm(get_text_width(document))
        run.add_picture(image_stream, width=width)
    except Exception as e:
        print(f"⚠ Warning: Could not add image: {e}")
        import traceback
        traceback.print_exc()
        # Continue without image

def replace(v: list | str, old_str: str, new_str: str):
    match v:
        case str():
            return v.replace(old_str, new_str)
        case list():
            return list(map(lambda iv: iv.replace(old_str, new_str), v))
        case _:
            return v

def generate_classplus_table_formatted_document(file_name, questions, folder_name):
    idx = 0
    document = Document()
    set_margins(document)
    for ques in questions:
        question = {k: replace(v, "*", "") for k, v in ques.items()}
        table = document.add_table(0, 3)
        table.style = 'TableGrid'
        row_cells = table.add_row().cells
        row_cells[0].text = "Question"
        paragraph = row_cells[1].paragraphs[0]
        question_before_table, rows, question_after_table = parse_question_and_table(question["question"])

        curr_style = add_text_with_style(paragraph, question_before_table.strip(), 0)
        if (len(rows) > 0):
            internal_table = row_cells[1].add_table(0, len(rows[0]))
            for row in rows:
                table_row = internal_table.add_row()
                table_row.height = Cm(1.1)
                cells = table_row.cells
                for idx, row_data in enumerate(row):
                    cells[idx].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    add_text_with_style(cells[idx].paragraphs[0], row_data, 0)

                internal_table.style = "Table Grid"
                internal_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        curr_style = add_text_with_style(row_cells[1].add_paragraph(), question_after_table.strip(), curr_style)
        row_cells = table.add_row().cells
        row_cells[0].text = "Type"
        row_cells[1].text = "multiple_choice"
        option_idx = 0
        for option in question["options"]:
            row_cells = table.add_row().cells
            row_cells[0].text = "Option"
            row_cells[1].text = option
            row_cells[2].text = "correct" if question["answer"] == option_idx else "incorrect"
            option_idx += 1

        row_cells = table.add_row().cells
        row_cells[0].text = "Solution"
        paragraph = row_cells[1].paragraphs[0]
        paragraph.add_run(f"Answer: ({chr(97 + question["answer"])}) {question["options"][question["answer"]]}" + "\nExplanation: ").bold = True
        add_text_with_style(paragraph, question["explanation"] + "\n", curr_style=0)

        if (len(question["source"]) > 0):
            paragraph.add_run(f"{question["source"]}").bold = True

        row_cells = table.add_row().cells
        row_cells[0].text = "Marks"
        row_cells[1].text = "2"
        row_cells[2].text = "0.67"
        
        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(1, 1).merge(table.cell(1, 2))
        table.cell(6, 1).merge(table.cell(6, 2))

        document.add_page_break()
        idx += 1

    file_name_without_ext = file_name.replace(".txt", "").replace(".md", "")
    file_path = os.path.join(folder_name, "output", f"converted_{file_name_without_ext}-classplus.docx")
    file_path = os.path.normpath(file_path)  # Normalize path for Windows
    
    # Auto-handle file saving with retries
    for retry in range(3):
        try:
            if retry > 0 and is_windows():
                close_word_if_needed()
                time.sleep(0.5)
            
            # Remove existing file if it exists
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    time.sleep(0.2)
                except PermissionError:
                    if retry < 2:
                        continue
                    if is_windows():
                        close_word_if_needed()
                        time.sleep(1)
                        try:
                            os.remove(file_path)
                        except:
                            pass
            
            # Post-process: Convert any remaining markdown tables to Word tables
            try:
                from table_converter import convert_markdown_tables_in_paragraphs
                convert_markdown_tables_in_paragraphs(document)
            except Exception as e:
                # If conversion fails, continue anyway
                pass
            
            document.save(file_path)
            if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                print(f"✓ Created {os.path.basename(file_path)}")
                return
        except PermissionError:
            if retry < 2:
                if is_windows():
                    close_word_if_needed()
                time.sleep(1)
                continue
            else:
                print(f"⚠ Skipped {os.path.basename(file_path)} (file locked)")
                return
        except Exception as e:
            if retry < 2:
                time.sleep(0.5)
                continue
            else:
                print(f"⚠ Skipped {os.path.basename(file_path)}: {str(e)[:50]}")
                return

def process_text_with_images(text: str, images: list, document, current_paragraph):
    """
    Process text and insert images at [IMAGE:N] markers
    Returns: (processed_text, remaining_images)
    Images are inserted into the document at marker positions
    """
    import re
    if not images:
        # Remove any image markers if no images available
        return re.sub(r'\[IMAGE:\d+\]\n?', '', text), []
    
    remaining_images = images.copy()
    lines = text.splitlines(keepends=True)
    processed_lines = []
    image_para_added = False
    
    for line in lines:
        # Check if line contains image marker
        match = re.search(r'\[IMAGE:(\d+)\]', line)
        if match:
            image_idx = int(match.group(1))
            if image_idx < len(remaining_images) and remaining_images[image_idx]:
                # Add image to document
                add_image(document, remaining_images[image_idx])
                image_para_added = True
                # Remove marker from line
                line = re.sub(r'\[IMAGE:\d+\]\n?', '', line)
                # Mark image as used (optional)
                # remaining_images[image_idx] = None
        
        # Add line if it has content (or if it's not just an image marker)
        if line.strip() or not re.search(r'\[IMAGE:\d+\]', line):
            processed_lines.append(line)
    
    processed_text = ''.join(processed_lines)
    return processed_text, remaining_images

def add_text_with_style(paragraph, text: str, curr_style):
    """
    Parse text with markdown-style formatting (**bold**, *italic*) and apply to paragraph
    curr_style: 0=normal, 1=bold, 2=italic, 3=bold+italic
    """
    if (len(text.strip()) == 0):
        return curr_style

    curr_text = ""
    texts = []
    i = 0
    
    while (i < len(text)):
        ch = text[i]
        
        # Check for **bold** (double asterisk)
        if (ch == "*" and i + 1 < len(text) and text[i + 1] == "*"):
            # Save current text with current style
            if (len(curr_text) > 0):
                texts.append({
                    "text": curr_text,
                    "style": curr_style
                })
                curr_text = ""
            
            # Toggle bold (0<->1, 2<->3)
            curr_style = curr_style ^ 1
            i += 2  # Skip both asterisks
            continue
        
        # Check for *italic* (single asterisk, but not part of **)
        elif (ch == "*" and (i == 0 or text[i-1] != "*") and (i + 1 >= len(text) or text[i+1] != "*")):
            # Save current text with current style
            if (len(curr_text) > 0):
                texts.append({
                    "text": curr_text,
                    "style": curr_style
                })
                curr_text = ""
            
            # Toggle italic (0<->2, 1<->3)
            if curr_style in [0, 1]:
                curr_style = curr_style ^ 2
            else:
                curr_style = curr_style ^ 2
            i += 1  # Skip asterisk
            continue
        else:
            curr_text += ch

        i += 1

    # Add remaining text
    if (len(curr_text) > 0):
        texts.append({
            "text": curr_text,
            "style": curr_style
        })

    # Apply formatting to paragraph
    for text_obj in texts:
        run = paragraph.add_run(text_obj["text"])
        
        # Apply style based on curr_style value
        style_val = text_obj["style"]
        if style_val == 0:  # Normal
            run.bold = False
            run.italic = False
        elif style_val == 1:  # Bold only
            run.bold = True
            run.italic = False
            # Explicitly ensure bold is set in XML for better PDF compatibility
            try:
                if hasattr(run, '_element'):
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        run._element.insert(0, rPr)
                    b_elem = rPr.find(qn('w:b'))
                    if b_elem is None:
                        b_elem = OxmlElement('w:b')
                        rPr.append(b_elem)
                    b_elem.set(qn('w:val'), 'true')
            except:
                pass  # If XML manipulation fails, run.bold = True should still work
        elif style_val == 2:  # Italic only
            run.bold = False
            run.italic = True
        elif style_val == 3:  # Bold + Italic
            run.bold = True
            run.italic = True
            # Explicitly ensure bold is set in XML for better PDF compatibility
            try:
                if hasattr(run, '_element'):
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        run._element.insert(0, rPr)
                    b_elem = rPr.find(qn('w:b'))
                    if b_elem is None:
                        b_elem = OxmlElement('w:b')
                        rPr.append(b_elem)
                    b_elem.set(qn('w:val'), 'true')
            except:
                pass  # If XML manipulation fails, run.bold = True should still work
            
    return 0  # Reset to normal style after processing

def o_add_text_with_style(paragraph, text):
    curr_text = ""
    texts = []
    curr_style = "n"
    i = 0

    while (i < len(text)):
        ch = text[i]
        if (ch == "<" and i + 1 < len(text) and text[i + 1] in BOLD_ITALIC_IDENTIFIER.keys() and i + 2 < len(text) and text[i + 2] == '>'):
            if (len(curr_text) > 0):
                texts.append({
                    "text": curr_text,
                    "style": curr_style
                })

                curr_text = ""
            
            curr_style = BOLD_ITALIC_IDENTIFIER[text[i + 1]]
            i += 2
            
        elif (ch == "<" and i + 1 < len(text) and text[i + 1] == "/" and i + 2 < len(text) and text[i + 2] in BOLD_ITALIC_IDENTIFIER.keys() and i + 3 < len(text) and text[i + 3] == ">"):
            if (len(curr_text) > 0):
                texts.append({
                    "text": curr_text,
                    "style": curr_style
                })

                curr_text = ""
            
            curr_style = "N"
            i += 3
        else:
            curr_text += ch

        i += 1

    if (len(curr_text) > 0):
        texts.append({
            "text": curr_text,
            "style": curr_style
        })

    for text_obj in texts:
        match text_obj["style"]:
            case "b":
                paragraph.add_run(text_obj["text"], style="NormalStyle").bold = True
            case "i":
                paragraph.add_run(text_obj["text"], style="NormalStyle").italic = True 
            case _:
                paragraph.add_run(text_obj["text"], style="NormalStyle")



def process_question_content_sequentially(content: str, images: list, document, ques_no: int, page_break_before: bool = False):
    """
    Process question content sequentially: text, images, tables in exact order
    Returns: (remaining_images, curr_style, deferred_images, deferred_difficulty)
    where deferred_images is a list of image indices that should be inserted after explanation
    and deferred_difficulty is the difficulty text to be added after explanation
    """
    import re
    lines = content.splitlines(keepends=True)
    curr_style = 0
    current_paragraph = None
    current_text = ""
    in_table = False
    table_rows = []
    table_started = False
    first_paragraph_added = False
    deferred_images = []  # Store image indices to insert after explanation
    deferred_difficulty = None  # Store difficulty text to add after explanation
    
    # Add question number at start
    if lines:
        current_paragraph = document.add_paragraph()
        current_paragraph.paragraph_format.line_spacing = line_spacing
        current_paragraph.paragraph_format.space_after = Pt(0)  # No space after question paragraphs
        if page_break_before:
            current_paragraph.paragraph_format.page_break_before = True
        # Start with question number - accumulate it
        current_text = bold(f"Q{ques_no}. ")
        first_paragraph_added = True
    
    for line in lines:
        stripped = line.strip()
        
        # CRITICAL: Check for joined table rows with || BEFORE other processing
        # This must happen first to prevent rows from being incorrectly parsed
        if stripped.startswith("|") and "||" in line:  # Check original line, not stripped
            # IMPORTANT: Add any accumulated text BEFORE processing the table
            # This ensures text like "Consider the following informations:" is added before the table
            if current_text.strip() and current_paragraph:
                curr_style = add_text_with_style(current_paragraph, current_text, curr_style)
                current_text = ""
            elif current_text.strip() and not current_paragraph:
                # Create paragraph if we have text but no paragraph yet
                current_paragraph = document.add_paragraph()
                current_paragraph.paragraph_format.line_spacing = line_spacing
                current_paragraph.paragraph_format.space_after = Pt(0)  # No space after question paragraphs
                if page_break_before and not first_paragraph_added:
                    current_paragraph.paragraph_format.page_break_before = True
                    first_paragraph_added = True
                curr_style = add_text_with_style(current_paragraph, current_text, curr_style)
                current_text = ""
            
            print(f"[DEBUG] EARLY DETECTION: Found joined table rows with ||")
            # Split by || to get individual rows
            joined_rows = line.split("||")
            for joined_row_str in joined_rows:
                joined_row_str = joined_row_str.strip()
                if not joined_row_str:
                    continue
                # Ensure it starts and ends with |
                if not joined_row_str.startswith("|"):
                    joined_row_str = "|" + joined_row_str
                if not joined_row_str.endswith("|"):
                    joined_row_str = joined_row_str + "|"
                
                # Parse this row
                parts = joined_row_str.split("|")
                column_values = [part.strip() for part in parts[1:-1]]
                
                # Skip separator rows
                is_separator = (
                    joined_row_str.startswith("|---") or 
                    (joined_row_str.startswith("|") and 
                     all(c in "-: |" for c in joined_row_str.replace(" ", "").replace("|", "")))
                )
                if is_separator:
                    print(f"[DEBUG] Skipping separator: {joined_row_str[:50]}")
                    continue
                
                # Check if separator row (cells are all dashes)
                if len(column_values) > 0:
                    is_sep_row = all(
                        all(c in "-: " for c in cell.strip()) 
                        for cell in column_values if cell.strip()
                    )
                    if not is_sep_row:
                        table_rows.append(column_values)
                        in_table = True
                        table_started = True
                        print(f"[DEBUG] Added row from || split: {len(column_values)} cells = {column_values[:2]}...")
            continue  # Skip normal processing for this line
        
        # Check for image marker in the line
        image_match = re.search(r'\[IMAGE:(\d+)\]', line)
        if image_match:
            # Split line into parts: before marker, marker, after marker
            before_marker = line[:image_match.start()]
            after_marker = line[image_match.end():]
            
            # CRITICAL: Add ALL accumulated text BEFORE inserting image
            # This ensures question text appears first, then image, then difficulty
            # Step 1: Add any accumulated text from previous lines to existing paragraph
            if current_text.strip():
                # We have accumulated text from previous lines - MUST add it now
                # If we have an existing paragraph, use it; otherwise create a new one
                if not current_paragraph:
                    current_paragraph = document.add_paragraph()
                    current_paragraph.paragraph_format.line_spacing = line_spacing
                    current_paragraph.paragraph_format.space_after = Pt(0)  # No space after question paragraphs
                    # If this is the first paragraph and we need a page break, set it
                    if page_break_before and not first_paragraph_added:
                        current_paragraph.paragraph_format.page_break_before = True
                        first_paragraph_added = True
                else:
                    # Existing paragraph - ensure no space after
                    current_paragraph.paragraph_format.space_after = Pt(0)
                # CRITICAL: Add accumulated text to the paragraph BEFORE image insertion
                # This ensures the text paragraph is created and added to document first
                curr_style = add_text_with_style(current_paragraph, current_text, curr_style)
                current_text = ""
                # Keep current_paragraph for now - we'll reset after image
            
            # Step 2: Add any text that appears before the marker on the same line
            if before_marker.strip():
                # If we still have a paragraph from step 1, use it; otherwise create new
                if not current_paragraph:
                    current_paragraph = document.add_paragraph()
                    current_paragraph.paragraph_format.line_spacing = line_spacing
                    current_paragraph.paragraph_format.space_after = Pt(0)  # No space after question paragraphs
                    if page_break_before and not first_paragraph_added:
                        current_paragraph.paragraph_format.page_break_before = True
                        first_paragraph_added = True
                curr_style = add_text_with_style(current_paragraph, before_marker, curr_style)
            
            # Step 3: Finalize text paragraph and ensure it's added to document
            # We need to make sure the text paragraph is completely finalized before inserting image
            text_paragraph_finalized = False
            if current_paragraph:
                # Paragraph already has text added via add_text_with_style
                # The paragraph is already in the document, we just need to reset our reference
                text_paragraph_finalized = True
                current_paragraph = None
            elif current_text.strip():
                # We have text but no paragraph - this shouldn't happen, but handle it
                current_paragraph = document.add_paragraph()
                current_paragraph.paragraph_format.line_spacing = line_spacing
                current_paragraph.paragraph_format.space_after = Pt(0)  # No space after question paragraphs
                if page_break_before and not first_paragraph_added:
                    current_paragraph.paragraph_format.page_break_before = True
                    first_paragraph_added = True
                curr_style = add_text_with_style(current_paragraph, current_text, curr_style)
                current_text = ""
                text_paragraph_finalized = True
                current_paragraph = None
            
            # Step 4: Defer image insertion until after explanation
            # Images should appear after the explanation, not in the question content
            image_idx = int(image_match.group(1))
            if image_idx < len(images) and images[image_idx]:
                print(f"  → [IMAGE DEFER] Deferring image {image_idx} for question {ques_no} (will insert after explanation)")
                # Store image index to insert after explanation
                deferred_images.append(image_idx)
            else:
                if image_idx >= len(images):
                    print(f"⚠ Warning: Image index {image_idx} out of range (only {len(images)} images available)")
                elif not images[image_idx]:
                    print(f"⚠ Warning: Image {image_idx} is None or empty")
            
            # Step 5: Process text after marker on the same line (e.g., if image marker is inline with text)
            if after_marker.strip():
                # Add text after image immediately to maintain correct order
                # Create new paragraph after image (image creates its own paragraph)
                current_paragraph = document.add_paragraph()
                current_paragraph.paragraph_format.line_spacing = line_spacing
                current_paragraph.paragraph_format.space_after = Pt(0)  # No space after question paragraphs
                curr_style = add_text_with_style(current_paragraph, after_marker, curr_style)
                current_text = ""  # Clear accumulated text since we just added it
            else:
                # No text after marker on same line
                # The next line (e.g., "Difficulty: Moderate") will be processed in the next iteration
                # CRITICAL: Reset paragraph so next line creates a new one AFTER the image
                # This ensures "Difficulty: Moderate" appears after the image, not before
                current_text = ""
                current_paragraph = None
            
            # Skip to next line (don't process this line again)
            continue
        
        # Check if line is a table separator first (before checking for table row)
        # Separator rows are like |---|---| or | --- | --- | --- |
        is_separator_line = (
            stripped.startswith("|---") or 
            (stripped.startswith("|") and 
             len(stripped) > 1 and
             all(c in "-: |" for c in stripped.replace(" ", "").replace("|", "")) and
             len([c for c in stripped if c == "|"]) >= 2)  # Has at least 2 pipe characters
        )
        if is_separator_line:
            # Table separator - skip it completely, don't add to table_rows
            continue
        
        # Check if line is a table row
        if stripped.startswith("|"):
            # Add any accumulated text before starting table
            if current_text.strip() and current_paragraph:
                curr_style = add_text_with_style(current_paragraph, current_text, curr_style)
                current_text = ""
                current_paragraph = None
            
            # CRITICAL FIX: Check if multiple table rows are joined with ||
            # If we see || in the line, it means multiple rows are joined
            if "||" in stripped:
                print(f"[DEBUG] Detected joined table rows with ||: {stripped[:100]}...")
                # Split by || to get individual rows
                joined_rows = stripped.split("||")
                for joined_row in joined_rows:
                    joined_row = joined_row.strip()
                    if not joined_row:
                        continue
                    # Ensure it starts and ends with |
                    if not joined_row.startswith("|"):
                        joined_row = "|" + joined_row
                    if not joined_row.endswith("|"):
                        joined_row = joined_row + "|"
                    
                    # Parse this row
                    parts = joined_row.split("|")
                    column_values = [part.strip() for part in parts[1:-1]]  # Preserve all cells including empty ones
                    
                    # Check for separator
                    is_separator_line = (
                        joined_row.startswith("|---") or 
                        (joined_row.startswith("|") and 
                         len(joined_row) > 1 and
                         all(c in "-: |" for c in joined_row.replace(" ", "").replace("|", "")) and
                         len([c for c in joined_row if c == "|"]) >= 2)
                    )
                    if is_separator_line:
                        print(f"[DEBUG] Skipping separator row from joined rows")
                        continue
                    
                    # Check if separator row
                    if len(column_values) > 0:
                        is_separator_row = all(
                            all(c in "-: " for c in cell.strip()) 
                            for cell in column_values if cell.strip()
                        )
                        if not is_separator_row:
                            table_rows.append(column_values)
                            in_table = True
                            table_started = True
                            print(f"[DEBUG] Added row from joined rows: {len(column_values)} cells")
                    continue
            
            # Parse table row - preserve empty cells
            # Split by | and take elements from index 1 to -1 (excluding first/last empty strings)
            parts = stripped.split("|")
            column_values = [part.strip() for part in parts[1:-1]]  # Preserve all cells including empty ones
            
            # DEBUG: Check for suspiciously long rows (might be multiple rows joined)
            if len(column_values) > 10:
                print(f"[DEBUG] WARNING: Table row has {len(column_values)} cells - might be multiple rows joined!")
                print(f"  Line content: {stripped[:100]}...")
                # Try to detect if this is actually multiple table rows joined (common column counts)
                for test_cols in [2, 3, 4, 5]:
                    if len(column_values) % test_cols == 0:
                        num_potential_rows = len(column_values) // test_cols
                        print(f"  [DEBUG] Could be {num_potential_rows} rows of {test_cols} columns")
                        # Split into potential rows
                        potential_rows = [column_values[j:j+test_cols] for j in range(0, len(column_values), test_cols)]
                        # Check if split makes sense (not all separator rows)
                        non_sep_count = sum(1 for row in potential_rows 
                                          if not all(all(c in "-: " for c in cell.strip()) 
                                                   for cell in row if cell.strip()))
                        if non_sep_count >= 2:  # At least 2 non-separator rows
                            print(f"  [DEBUG] Splitting joined row into {len(potential_rows)} separate rows")
                            # Add each as a separate row (AGGRESSIVELY filter out separators)
                            for potential_row in potential_rows:
                                # AGGRESSIVE separator detection
                                is_sep = True
                                has_content = False
                                for cell in potential_row:
                                    cell_stripped = cell.strip()
                                    if cell_stripped:
                                        has_content = True
                                        # Check if cell contains only separator characters
                                        remaining = cell_stripped.replace("-", "").replace(":", "").replace("=", "").replace(" ", "").replace("|", "")
                                        if remaining:  # Has non-separator content
                                            is_sep = False
                                            break
                                
                                # Only add if it's not a separator row AND has content
                                if not is_sep and has_content:
                                    table_rows.append(potential_row)
                                    in_table = True
                                    table_started = True
                                    print(f"  [DEBUG] Added split row {len(table_rows)}: {potential_row}")
                                else:
                                    print(f"  [DEBUG] FILTERED OUT separator/empty row: {potential_row}")
                            continue
            
            # AGGRESSIVE separator detection: don't add separator rows
            if len(column_values) > 0:
                # Check if this is a separator row - more aggressive check
                is_separator_row = True
                has_content = False
                for cell in column_values:
                    cell_stripped = cell.strip()
                    if cell_stripped:
                        has_content = True
                        # Remove all separator characters and check if anything remains
                        remaining = cell_stripped.replace("-", "").replace(":", "").replace("=", "").replace(" ", "").replace("|", "")
                        if remaining:  # Has non-separator content
                            is_separator_row = False
                            break
                
                # Only add if it's NOT a separator row AND has content
                if not is_separator_row and has_content:
                    # CRITICAL: Ensure we're appending a list, not extending
                    # Double-check that column_values is a list
                    if isinstance(column_values, list):
                        table_rows.append(column_values)  # APPEND, not extend!
                        print(f"[DEBUG] Added table row {len(table_rows)}: {len(column_values)} cells = {column_values[:2]}...")
                    else:
                        print(f"[DEBUG] ERROR: column_values is not a list: {type(column_values)}")
                    in_table = True
                    table_started = True
                else:
                    print(f"[DEBUG] FILTERED OUT separator/empty row: {column_values}")
            continue
        
        # Regular text line (including text after image markers)
        else:
            # If we were in a table, close it first
            if in_table and table_rows:
                # Validate table structure before closing
                print(f"[DEBUG] Closing table with {len(table_rows)} rows")
                for i, row in enumerate(table_rows):
                    print(f"  [DEBUG] Table row {i}: {len(row)} cells")
                    if len(row) > 10:
                        print(f"    [DEBUG] WARNING: Row {i} has {len(row)} cells - might be flattened!")
                
                # Filter out separator rows - be very aggressive about this
                filtered_rows = []
                for row in table_rows:
                    if not row:  # Skip empty rows
                        continue
                    # Check if all non-empty cells contain only dashes, colons, or spaces
                    is_separator = all(
                        all(c in "-: " for c in cell.strip()) 
                        for cell in row if cell.strip()
                    )
                    if not is_separator:
                        filtered_rows.append(row)
                
                if len(filtered_rows) > 0:
                    # Add spacing before table
                    table_para = document.add_paragraph()
                    table_para.paragraph_format.space_before = Pt(6)
                    add_table(document, filtered_rows)
                    # Add spacing after table
                    table_para_after = document.add_paragraph()
                    table_para_after.paragraph_format.space_after = Pt(6)
                
                table_rows = []
                in_table = False
                current_paragraph = None
                current_text = ""
            
            # Check if this is a difficulty line - defer it to appear after explanation
            if stripped.lower().startswith("difficulty:"):
                # Store difficulty text to be added after explanation
                deferred_difficulty = stripped.strip()
                print(f"  → [DIFFICULTY DEFER] Deferring difficulty text: {deferred_difficulty}")
                continue  # Skip processing this line
            
            # Add text to current paragraph (this includes text after image markers)
            if stripped or line:  # Include even if just whitespace/newline to preserve structure
                if not current_paragraph:
                    # Create new paragraph (e.g., after an image or at start)
                    current_paragraph = document.add_paragraph()
                    current_paragraph.paragraph_format.line_spacing = line_spacing
                    current_paragraph.paragraph_format.space_after = Pt(0)  # No space after question paragraphs
                    # If we just created a paragraph after an image, add text immediately
                    # to ensure correct order (text appears right after image)
                    if line.strip():
                        curr_style = add_text_with_style(current_paragraph, line, curr_style)
                        current_text = ""  # Clear since we just added it
                    else:
                        current_text += line  # Accumulate whitespace-only lines
                else:
                    # Accumulate text in existing paragraph
                    current_text += line
    
    # Handle remaining text (but check for difficulty text first)
    if current_text.strip():
        # Check if remaining text contains difficulty
        remaining_lines = current_text.splitlines()
        filtered_lines = []
        for remaining_line in remaining_lines:
            stripped_remaining = remaining_line.strip()
            if stripped_remaining.lower().startswith("difficulty:"):
                # Store difficulty text to be added after explanation
                if not deferred_difficulty:  # Only store if we haven't already found one
                    deferred_difficulty = stripped_remaining
                    print(f"  → [DIFFICULTY DEFER] Deferring difficulty from remaining text: {deferred_difficulty}")
                continue  # Skip this line
            filtered_lines.append(remaining_line)
        
        # Add remaining text (excluding difficulty)
        if filtered_lines:
            remaining_text = "\n".join(filtered_lines)
            if remaining_text.strip():
                if not current_paragraph:
                    current_paragraph = document.add_paragraph()
                    current_paragraph.paragraph_format.line_spacing = line_spacing
                    current_paragraph.paragraph_format.space_after = Pt(0)  # No space after question paragraphs
                curr_style = add_text_with_style(current_paragraph, remaining_text, curr_style)
    
    # Handle remaining table
    if in_table and table_rows:
        # Filter out separator rows - be VERY aggressive about this
        filtered_rows = []
        for row_idx, row in enumerate(table_rows):
            if not row:  # Skip empty rows
                print(f"[DEBUG] Skipping empty row {row_idx}")
                continue
            
            # AGGRESSIVE separator detection
            is_separator = True
            has_content = False
            for cell in row:
                cell_stripped = str(cell).strip() if cell else ""
                if cell_stripped:
                    has_content = True
                    # Remove all separator characters and check if anything remains
                    remaining = cell_stripped.replace("-", "").replace(":", "").replace("=", "").replace(" ", "").replace("|", "")
                    if remaining:  # Has non-separator content
                        is_separator = False
                        break
            
            # Only add if it's NOT a separator row AND has content
            if has_content and not is_separator:
                filtered_rows.append(row)
                print(f"[DEBUG] Keeping row {row_idx} for table: {row[:2]}...")
            else:
                print(f"[DEBUG] FILTERED OUT row {row_idx} (separator or empty): {row}")
        
        if len(filtered_rows) > 0:
            table_para = document.add_paragraph()
            table_para.paragraph_format.space_before = Pt(6)
            add_table(document, filtered_rows)
            table_para_after = document.add_paragraph()
            table_para_after.paragraph_format.space_after = Pt(6)
    
    return images, curr_style, deferred_images, deferred_difficulty

def add_questions_and_explanation(questions, document, images, page_break_before = True):
    answers = [-1] + list(map(lambda ques: f"{chr(ord("a") + ques["answer"])}) {ques["options"][ques["answer"]]}", questions))
    ques_no = 0
    curr_style = 0
    for question in questions:
        ques_no += 1
        ques, options = question["question"], question["options"]
        
        # Set page break for first question
        if page_break_before and ques_no == 1:
            # Will be set on first paragraph
            pass
        elif not page_break_before:
            page_break_before = True
        
        # Process question content sequentially (text, images, tables in order)
        should_page_break = page_break_before and ques_no == 1
        images, curr_style, deferred_images, deferred_difficulty = process_question_content_sequentially(ques, images, document, ques_no, should_page_break)
        
        # Remove extra space after last question paragraph (before options)
        # Find the last paragraph and set space_after to 0
        if len(document.paragraphs) > 0:
            last_para = document.paragraphs[-1]
            last_para.paragraph_format.space_after = Pt(0)
        
        # Add options with formatting preserved
        options_idx = ["a) ", "b) ", "c) ", "d) "]
        for idx, option in enumerate(zip(options_idx, options)):
            o_paragraph = document.add_paragraph()
            # Only add space before if it's not the first option
            if idx > 0:
                o_paragraph.paragraph_format.space_before = Pt(1.2)
            else:
                o_paragraph.paragraph_format.space_before = Pt(0)  # No space before first option
            o_paragraph.paragraph_format.space_after = Pt(0)  # No space after options
            curr_style = add_text_with_style(o_paragraph, option[0] + option[1].strip(), curr_style)
            
        ae_paragraph = document.add_paragraph()
        ae_paragraph.paragraph_format.space_before = Cm(1.1)
        ae_paragraph.paragraph_format.line_spacing = line_spacing

        content = "\n".join([bold(f"Answer: {answers[ques_no]}"), bold("Explanation:"), question["explanation"]])
        curr_style = add_text_with_style(ae_paragraph, content, curr_style)
        
        # Insert deferred images AFTER explanation
        for image_idx in deferred_images:
            if image_idx < len(images) and images[image_idx]:
                try:
                    print(f"  → Inserting deferred image {image_idx} after explanation for question {ques_no}")
                    add_image(document, images[image_idx])
                    print(f"  ✓ Deferred image {image_idx} inserted successfully after explanation")
                except Exception as e:
                    print(f"⚠ Warning: Could not insert deferred image {image_idx}: {e}")
                    import traceback
                    traceback.print_exc()
        
        # Add deferred difficulty text AFTER images
        if deferred_difficulty:
            difficulty_paragraph = document.add_paragraph()
            difficulty_paragraph.paragraph_format.space_before = Pt(6)
            difficulty_paragraph.paragraph_format.line_spacing = line_spacing
            curr_style = add_text_with_style(difficulty_paragraph, deferred_difficulty, curr_style)
            print(f"  ✓ Deferred difficulty text added after explanation: {deferred_difficulty}")
        
        # Note: Images are already inserted in their correct positions during sequential processing
        # Do NOT add images again here - they're already in the question content

        sd_paragraph = document.add_paragraph()
        sd_paragraph.paragraph_format.space_before = Cm(0.8)
        sd_paragraph.paragraph_format.line_spacing = line_spacing
        curr_style = add_text_with_style(sd_paragraph, bold(question["source"]), curr_style)

def add_answer_key(questions, document: Document, add_page_break = False):
    if (add_page_break):
        document.add_page_break()

    heading = document.add_paragraph()
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("ANSWERS", style="NormalStyle")
    run.bold = True
    run.underline = True
    
    add_empty_paragraph(document)
    table = document.add_table(0, 10)
    answers = list(map(lambda ques: chr(ord("A") + ques["answer"]), questions))
    ques_no = 0
    while (len(answers) > 0):
        first_5_answers = answers[:5]
        cells = table.add_row().cells
        idx = 0
        for ans in first_5_answers:
            ques_no += 1
            cells[idx].width = Cm(1.7)
            cells[idx].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[idx].paragraphs[0].add_run(f"Q.{ques_no})", style="TableStyle").bold = True

            cells[idx + 1].width = Cm(1.7)
            cells[idx + 1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[idx + 1].paragraphs[0].add_run(ans.upper(), style="TableStyle").bold = True
            
            idx += 2

        answers = answers[len(first_5_answers):]

    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.height = Cm(1.2)
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_idx % 2 == 1:
                set_cell_background(cell, 'FEF3CC')

    return document

def is_not_empty(row: list[str]):
    is_empty = True
    for v in row:
        is_empty = is_empty and set(v.strip()).issubset(set("-:_="))

    return (not is_empty)

def parse_question_and_table(ques: str):
    """Parse question text and extract tables (markdown format)"""
    lines = ques.splitlines()
    question_before = []
    question_after = []
    rows = []
    current_section = "before"  # "before", "table", "after"
    
    for line in lines:
        stripped_line = line.strip()
        
        # Check if line is a table separator (|---|)
        if stripped_line.startswith("|---") or (stripped_line.startswith("|") and all(c in "-: " for c in stripped_line.replace("|", "").strip())):
            # Table separator, mark that we're in a table
            if current_section == "before":
                current_section = "table"
            continue
        # Check if line is a table row (starts with |)
        elif stripped_line.startswith("|"):
            # It's a table row
            # Parse table row - preserve empty cells to maintain column alignment
            # Split by | and take elements from index 1 to -1 (excluding first/last empty strings)
            parts = stripped_line.split("|")
            columnValues = [part.strip() for part in parts[1:-1]]  # Preserve all cells including empty ones
            if len(columnValues) > 0:
                rows.append(columnValues)
                if current_section == "before":
                    current_section = "table"
        else:
            # Regular text line
            if current_section == "table":
                # We've left the table section, start collecting after-table text
                current_section = "after"
                question_after.append(line)
            elif current_section == "before":
                question_before.append(line)
            else:  # current_section == "after"
                question_after.append(line)

    # Filter out empty rows and separator rows
    filtered_rows = []
    for row in rows:
        if is_not_empty(row):
            # Check if it's a separator row
            is_separator = all(
                all(c in "-: " for c in cell.strip()) 
                for cell in row if cell.strip()
            )
            if not is_separator:
                filtered_rows.append(row)
    
    # Join question parts
    question_before_text = "\n".join(question_before).strip()
    question_after_text = "\n".join(question_after).strip()
    
    return question_before_text, filtered_rows, question_after_text

def add_table(document, rows: list[list[str]]):
    """Add table to document with formatting preserved"""
    if (rows == None or len(rows) == 0):
        return

    if len(rows) == 0:
        return
    
    # DEBUG: Print what we receive
    print(f"\n[DEBUG] add_table called with {len(rows)} rows")
    for i, row in enumerate(rows):
        print(f"  [DEBUG] Row {i}: {row} (length: {len(row)})")
        if isinstance(row, list) and len(row) > 10:
            print(f"    [DEBUG] WARNING: Row {i} has {len(row)} cells - might be flattened!")
            # If a row has more than 10 cells, it's likely flattened - try to split it
            # This is a workaround for a bug where rows are being combined
            
            # Try to determine column count by checking other rows or common patterns
            expected_cols = None
            # First, check if other rows give us a clue
            for other_row in rows:
                if isinstance(other_row, list) and 2 <= len(other_row) <= 10:
                    expected_cols = len(other_row)
                    break
            
            # If no clue from other rows, try common column counts (2, 3, 4, 5)
            if expected_cols is None:
                for test_cols in [2, 3, 4, 5]:
                    if len(row) % test_cols == 0:
                        expected_cols = test_cols
                        break
            
            # If we found a likely column count, split the row
            if expected_cols and len(row) % expected_cols == 0:
                num_split_rows = len(row) // expected_cols
                print(f"    [DEBUG] Attempting to split flattened row into {num_split_rows} rows of {expected_cols} columns")
                # Split into chunks
                split_rows = [row[j:j+expected_cols] for j in range(0, len(row), expected_cols)]
                # Filter out separator rows from split
                filtered_split = []
                for split_row in split_rows:
                    is_sep = all(all(c in "-: " for c in cell.strip()) for cell in split_row if cell.strip())
                    if not is_sep:
                        filtered_split.append(split_row)
                if len(filtered_split) > 0:
                    print(f"    [DEBUG] Split into {len(filtered_split)} rows")
                    rows = filtered_split + [r for r in rows[i+1:] if isinstance(r, list) and len(r) <= 10]
                    break
            else:
                print(f"    [DEBUG] Could not determine column count for splitting. Row length: {len(row)}")
                # If this is the only row and it's clearly flattened, try common column counts
                if len(rows) == 1 and len(row) > 10:
                    print(f"    [DEBUG] Single flattened row detected, trying common column counts...")
                    for test_cols in [2, 3, 4, 5, 6]:
                        if len(row) % test_cols == 0:
                            num_split_rows = len(row) // test_cols
                            print(f"    [DEBUG] Trying {test_cols} columns -> {num_split_rows} rows")
                            split_rows = [row[j:j+test_cols] for j in range(0, len(row), test_cols)]
                            # Check if split makes sense (not all separator rows)
                            non_sep_count = sum(1 for split_row in split_rows 
                                              if not all(all(c in "-: " for c in cell.strip()) 
                                                       for cell in split_row if cell.strip()))
                            if non_sep_count >= 2:  # At least 2 non-separator rows
                                print(f"    [DEBUG] Split successful: {len(split_rows)} rows")
                                rows = split_rows
                    break
    
    # Filter out separator rows FIRST before determining columns
    # Separator rows have cells that are all dashes, colons, or spaces
    filtered_rows = []
    for row in rows:
        if not row:  # Skip empty rows
            continue
        # Ensure row is a list, not a flattened structure
        if not isinstance(row, list):
            continue
        # Check if all non-empty cells contain only dashes, colons, or spaces
        is_separator = all(
            all(c in "-: " for c in cell.strip()) 
            for cell in row if cell.strip()
        )
        if not is_separator:
            filtered_rows.append(row)
        else:
            print(f"  [DEBUG] Filtered separator row: {row}")
    
    print(f"[DEBUG] After filtering: {len(filtered_rows)} rows")
    if len(filtered_rows) == 0:
        return
    
    # CRITICAL FIX: Check if any row has suspiciously many cells (likely flattened)
    # If we have rows with different cell counts, and one has way more, it's likely flattened
    row_lengths = [len(row) for row in filtered_rows]
    max_len = max(row_lengths) if row_lengths else 0
    min_len = min(row_lengths) if row_lengths else 0
    
    # SIMPLIFIED FIX: Only fix if we have a clear problem
    # If we have multiple rows with correct column count, don't modify them
    # Only fix if we have a single row with many cells (flattened)
    # OR if we have rows with mismatched lengths where one is clearly flattened
    should_fix = False
    if len(filtered_rows) == 1 and max_len > 10:
        # Single row with many cells - definitely needs fixing
        should_fix = True
    elif len(filtered_rows) > 1 and max_len > 10 and min_len <= 10 and max_len != min_len:
        # Multiple rows but one is clearly flattened (much longer than others)
        should_fix = True
    
    if should_fix:
        print(f"[DEBUG] Detected row(s) with {max_len} cells - likely flattened!")
        
        # Try to determine the correct column count
        expected_cols = None
        
        # First, check if rows have different lengths (one might be correct)
        if max_len != min_len:
            # Find the most common row length (excluding the suspiciously long one)
            from collections import Counter
            length_counts = Counter(row_lengths)
            # Remove the max length from consideration
            if max_len in length_counts:
                del length_counts[max_len]
            
            if length_counts:
                # Use the most common length as the expected column count
                expected_cols = length_counts.most_common(1)[0][0]
                print(f"[DEBUG] Expected column count based on other rows: {expected_cols}")
        
        # If we couldn't determine from other rows, try common column counts
        if expected_cols is None:
            # Try common table column counts: 3, 4, 2, 5 (prioritize 3 as most common)
            for test_cols in [3, 4, 2, 5]:
                # Even if not perfectly divisible, if we have a single row, try splitting anyway
                if len(filtered_rows) == 1:
                    # For single row, always try to split (even if not perfectly divisible)
                    num_test_rows = max_len // test_cols
                    if num_test_rows >= 2 and num_test_rows <= 20:
                        expected_cols = test_cols
                        print(f"[DEBUG] Determined column count: {test_cols} (would create ~{num_test_rows} rows from single flattened row)")
                        break
                elif max_len % test_cols == 0:
                    # Check if splitting makes sense
                    num_test_rows = max_len // test_cols
                    # If splitting gives us a reasonable number of rows
                    if num_test_rows >= 2 and num_test_rows <= 20:
                        expected_cols = test_cols
                        print(f"[DEBUG] Determined column count: {test_cols} (would create {num_test_rows} rows)")
                        break
        
        # If we still don't have expected_cols, default to 3 (most common table structure)
        if expected_cols is None:
            expected_cols = 3
            print(f"[DEBUG] Using default column count: 3 (forcing split even if not perfectly divisible)")
        
        # Now fix all rows
        if expected_cols and expected_cols < max_len:
            print(f"[DEBUG] Fixing rows to have {expected_cols} columns...")
            fixed_rows = []
            for row_idx, row in enumerate(filtered_rows):
                print(f"[DEBUG] Processing row {row_idx} for fixing: {len(row)} cells, first cell: '{row[0] if row else 'empty'}'")
                if len(row) == expected_cols:
                    # Row is correct, keep it
                    print(f"[DEBUG] Row {row_idx} is correct, keeping as-is")
                    fixed_rows.append(row)
                elif len(row) > expected_cols:
                    # Row is flattened, split it
                    # Calculate how many complete rows we can make
                    num_complete_rows = len(row) // expected_cols
                    remainder = len(row) % expected_cols
                    
                    if num_complete_rows > 0:
                        print(f"[DEBUG] Splitting row {row_idx} with {len(row)} cells into {num_complete_rows} rows of {expected_cols} columns (remainder: {remainder})")
                        print(f"[DEBUG] Row content: {row[:6]}...")
                        
                        # CRITICAL: Remove empty strings that are artifacts from || joins
                        # Empty strings at boundaries create misalignment
                        cleaned_row = []
                        for cell in row:
                            cell_str = str(cell).strip() if cell else ""
                            # Only skip if it's completely empty (not separator chars)
                            if cell_str:  # Keep non-empty cells
                                cleaned_row.append(cell_str)
                            # Skip empty strings - they're artifacts
                        
                        # Recalculate after cleaning
                        if len(cleaned_row) % expected_cols == 0:
                            num_complete_rows = len(cleaned_row) // expected_cols
                            print(f"[DEBUG] After cleaning empty cells: {len(cleaned_row)} cells, {num_complete_rows} rows")
                            row = cleaned_row
                        
                        # Split into complete rows - PRESERVE ORDER
                        # CRITICAL: Skip separator rows completely - don't add them at all
                        for j in range(0, num_complete_rows * expected_cols, expected_cols):
                            split_row = row[j:j+expected_cols]
                            print(f"[DEBUG]  Split chunk {j//expected_cols}: {split_row}")
                            
                            # AGGRESSIVE separator detection - check if ALL cells are separator-like
                            is_sep = True
                            has_content = False
                            for cell in split_row:
                                cell_stripped = cell.strip()
                                if cell_stripped:
                                    has_content = True
                                    # Check if cell contains only separator characters
                                    remaining = cell_stripped.replace("-", "").replace(":", "").replace("=", "").replace(" ", "").replace("|", "")
                                    if remaining:  # Has non-separator content
                                        is_sep = False
                                        break
                            
                            # Only add if it's not a separator row AND has content
                            if not is_sep and has_content:
                                fixed_rows.append(split_row)
                                print(f"[DEBUG]  Added as row {len(fixed_rows)-1}")
                            else:
                                print(f"[DEBUG]  FILTERED OUT: separator row or empty row")
                        
                        # DON'T add remainder rows - they're likely incomplete or separator cells
                        # If there's a remainder, it's probably separator cells or incomplete data
                        if remainder > 0:
                            print(f"[DEBUG] Ignoring {remainder} remainder cells (likely separator or incomplete)")
                    else:
                        # Row is shorter than expected_cols but we're here because len(row) > expected_cols
                        # This shouldn't happen, but handle it anyway
                        print(f"[DEBUG] Row {row_idx} {len(row)} cells, truncating to {expected_cols}")
                        fixed_rows.append(row[:expected_cols])
                else:
                    # Row has fewer cells, pad it
                    print(f"[DEBUG] Padding row {row_idx} from {len(row)} to {expected_cols} cells")
                    padded_row = row + [''] * (expected_cols - len(row))
                    fixed_rows.append(padded_row)
            
            filtered_rows = fixed_rows
            print(f"[DEBUG] After aggressive fix: {len(filtered_rows)} rows")
    
    # Determine number of columns (use max columns from filtered rows)
    # But be smart about it - if most rows have the same length, use that
    if not filtered_rows:
        return
    
    row_lengths = [len(row) for row in filtered_rows]
    from collections import Counter
    length_counts = Counter(row_lengths)
    
    # Use the most common row length as num_cols (more reliable than max)
    if length_counts:
        most_common_length, count = length_counts.most_common(1)[0]
        # If most rows have the same length, use that
        if count >= len(filtered_rows) * 0.5:  # At least 50% of rows have this length
            num_cols = most_common_length
            print(f"[DEBUG] Using most common row length: {num_cols} (appears in {count}/{len(filtered_rows)} rows)")
        else:
            # Use max if no clear majority
            num_cols = max(row_lengths)
            print(f"[DEBUG] Using max row length: {num_cols} (no clear majority)")
    else:
        num_cols = max(row_lengths) if row_lengths else 0
    
    if num_cols == 0:
        return
    
    # CRITICAL: Ensure all rows have the same number of columns
    # But preserve the exact order and data - only normalize column count
    # Also filter out rows that are completely empty
    normalized_rows = []
    for row_idx, row in enumerate(filtered_rows):
        # Skip completely empty rows
        if not row or all(not str(cell).strip() for cell in row):
            print(f"[DEBUG] Skipping completely empty row {row_idx}")
            continue
            
        if len(row) == num_cols:
            # Row is already correct, keep it exactly as is
            # But check if it's actually empty (all cells are empty)
            if any(str(cell).strip() for cell in row):
                normalized_rows.append(row)
                print(f"[DEBUG] Row {row_idx} normalized: {len(row)} cells (correct)")
            else:
                print(f"[DEBUG] Skipping row {row_idx}: all cells are empty")
        elif len(row) > num_cols:
            # Truncate if too many - but this shouldn't happen if fix worked correctly
            print(f"[DEBUG] WARNING: Row {row_idx} has {len(row)} cells, truncating to {num_cols}")
            print(f"[DEBUG]  Row data: {row}")
            truncated = row[:num_cols]
            # Only add if it has at least one non-empty cell
            if any(str(cell).strip() for cell in truncated):
                normalized_rows.append(truncated)
            else:
                print(f"[DEBUG] Skipping truncated row {row_idx}: all cells empty after truncation")
        else:
            # Row has fewer cells than expected
            # Since source table has no empty cells, this shouldn't happen
            # But if it does, check if this is a parsing error or real data
            print(f"[DEBUG] WARNING: Row {row_idx} has {len(row)} cells, expected {num_cols}")
            print(f"[DEBUG]  Row content: {row}")
            
            # Check if all other rows have the same length as this row
            other_row_lengths = [len(r) for r in filtered_rows if r != row]
            if other_row_lengths and len(set(other_row_lengths)) == 1 and other_row_lengths[0] == len(row):
                # All other rows have the same length as this row - num_cols was wrong
                print(f"[DEBUG]  All other rows have {other_row_lengths[0]} cells - num_cols was incorrectly set to {num_cols}")
                print(f"[DEBUG]  Recalculating num_cols...")
                # Recalculate num_cols based on actual row lengths
                num_cols = len(row)
                # Re-normalize all previously processed rows
                normalized_rows = [r[:num_cols] if len(r) >= num_cols else r for r in normalized_rows]
                normalized_rows.append(row)
            elif any(str(cell).strip() for cell in row):
                # This row is genuinely shorter - but don't pad with empty cells
                # Instead, just use what we have (this might indicate a parsing issue)
                print(f"[DEBUG]  Row is shorter but has content - using as-is (may indicate parsing issue)")
                normalized_rows.append(row)
            else:
                print(f"[DEBUG] Skipping row {row_idx}: all cells empty")
    
    filtered_rows = normalized_rows
    print(f"[DEBUG] Final table structure: {len(filtered_rows)} rows × {num_cols} columns")
    for i, r in enumerate(filtered_rows):
        print(f"  [DEBUG] Final row {i}: {r}")

    # Create table with proper structure
    table = document.add_table(rows=0, cols=num_cols)
    # Use None style to match input file (no predefined style)
    table.style = None
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False  # We'll set widths manually
    
    # Set table properties (width and borders)
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Add table width to ensure proper rendering
    tblW = OxmlElement('w:tblW')
    table_width = Inches(6.5)  # Match available width
    tblW.set(qn('w:w'), str(int(table_width)))  # Table width in EMUs
    tblW.set(qn('w:type'), 'dxa')  # Width type: dxa (twips)
    tblPr.append(tblW)
    
    # Add table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Set column widths - use full page width
    # Calculate available width (page width minus margins)
    # Standard page width is 8.5 inches, with 1 inch margins = 6.5 inches available
    # Use full available width to ensure columns are wide enough
    available_width_inches = 6.5  # Use full available width
    col_width_inches = available_width_inches / num_cols  # Distribute width evenly across columns
    col_width = Inches(col_width_inches)  # Convert to Inches object
    
    # Set column widths on table columns (more efficient and correct)
    for col in table.columns:
        col.width = col_width
    
    # DEBUG: Print all rows before processing - CRITICAL for debugging
    print(f"\n[DEBUG] ===== add_table called with {len(filtered_rows)} rows =====")
    for idx, r in enumerate(filtered_rows):
        print(f"  [DEBUG] Input row {idx}: {len(r)} cells = {r}")
    print(f"[DEBUG] ===========================================\n")
    
    # FINAL FILTER: Remove any separator rows that might have slipped through
    final_rows = []
    for row_idx, row in enumerate(filtered_rows):
        # AGGRESSIVE separator detection
        is_separator = True
        has_content = False
        
        for cell in row:
            cell_stripped = str(cell).strip() if cell else ""
            if cell_stripped:
                has_content = True
                # Remove all separator characters and check if anything remains
                remaining = cell_stripped.replace("-", "").replace(":", "").replace("=", "").replace(" ", "").replace("|", "")
                if remaining:  # Has non-separator content
                    is_separator = False
                    break
        
        # Only keep rows that have content and are not separators
        if has_content and not is_separator:
            final_rows.append(row)
            print(f"[DEBUG] Keeping row {row_idx}: {row[:2]}...")
        else:
            print(f"[DEBUG] FINAL FILTER: Removing row {row_idx} (separator or empty): {row}")
    
    filtered_rows = final_rows
    print(f"[DEBUG] After final filter: {len(filtered_rows)} rows")
    
    # ONE MORE CHECK: Remove any rows that are completely empty (all cells empty or whitespace)
    truly_final_rows = []
    for row_idx, row in enumerate(filtered_rows):
        # Check if row has any non-empty content
        has_any_content = False
        for cell in row:
            if cell and str(cell).strip():
                has_any_content = True
                break
        
        if has_any_content:
            truly_final_rows.append(row)
            print(f"[DEBUG] Keeping row {row_idx} (has content): {row[:2]}...")
        else:
            print(f"[DEBUG] REMOVING completely empty row {row_idx}: {row}")
    
    filtered_rows = truly_final_rows
    print(f"[DEBUG] After removing empty rows: {len(filtered_rows)} rows")
    
    if len(filtered_rows) == 0:
        print("[DEBUG] No rows to add to table after filtering!")
        return
    
    # Track actual table row number (separate from rows list index)
    # CRITICAL: Process rows in EXACT order they appear - do not reorder!
    # All separator rows and empty rows have been filtered out above
    actual_row_num = 0
    for row_idx, row in enumerate(filtered_rows):
        print(f"[DEBUG] Adding row {row_idx} to table (row {actual_row_num}): first cell='{row[0] if row else 'empty'}'")
            
        table_row = table.add_row()
        cells = table_row.cells
        
        # Set cell properties for better appearance
        for cell in cells:
            # Set vertical alignment to center for better appearance
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Add some padding for readability
            cell_para = cell.paragraphs[0]
            cell_para.paragraph_format.space_before = Pt(6)
            cell_para.paragraph_format.space_after = Pt(6)
        
        # Fill cells with data (column widths already set on table.columns)
        print(f"[DEBUG] Adding row {actual_row_num} to table with {len(row)} cells: {row[:3]}...")
        
        # Check if this row has any non-empty content BEFORE creating cells
        row_has_content = False
        for cell_data in row:
            if cell_data and str(cell_data).strip():
                cell_clean = str(cell_data).strip().replace("-", "").replace(":", "").replace("=", "").replace(" ", "").replace("|", "")
                if cell_clean:  # Has actual content (not just separator chars)
                    row_has_content = True
                    break
        
        if not row_has_content:
            print(f"[DEBUG] SKIPPING row {actual_row_num}: all cells are empty or separator-only")
            # Remove the row we just added
            try:
                table._element.remove(table_row._element)
            except:
                pass
            continue
        
        for idx in range(num_cols):
            if idx < len(row):
                row_data = str(row[idx]).strip() if row[idx] else ""
                # Remove formatting markers for cell content but preserve structure
                cell_para = cells[idx].paragraphs[0]
                # Clear any existing content
                cell_para.clear()
                # Use LEFT alignment for clean appearance
                cell_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Set padding for readability
                cell_para.paragraph_format.space_before = Pt(6)
                cell_para.paragraph_format.space_after = Pt(6)
                # FILTER OUT separator cells and empty cells
                cell_stripped_clean = row_data.replace("-", "").replace(":", "").replace("=", "").replace(" ", "").replace("|", "")
                is_separator_cell = (not cell_stripped_clean and row_data.strip())  # Empty after removing separators but had content
                
                # Use add_text_with_style to preserve bold/italic in cells
                if row_data and not is_separator_cell:
                    # Make first row (header) bold - use actual_row_num to track first table row
                    if actual_row_num == 0:
                        # Header row - make text bold
                        run = cell_para.add_run(row_data)
                        run.bold = True
                        run.font.size = Pt(11)  # Slightly larger for header
                        print(f"  [DEBUG] Cell {idx}: '{row_data}' (header, bold)")
                    else:
                        # Data rows - preserve formatting from markdown
                        add_text_with_style(cell_para, row_data, 0)
                        print(f"  [DEBUG] Cell {idx}: '{row_data}' (data)")
                else:
                    # Empty or separator cell - leave it empty (don't add any content)
                    cell_para.clear()
                    if is_separator_cell:
                        print(f"  [DEBUG] Cell {idx}: FILTERED OUT separator cell '{row_data}'")
                    else:
                        print(f"  [DEBUG] Cell {idx}: empty (no content)")
            else:
                # Cell index beyond row length - leave empty
                cells[idx].paragraphs[0].clear()
                print(f"  [DEBUG] Cell {idx}: empty (beyond row length)")
        
        # Increment actual row number after adding row to table
        actual_row_num += 1
        print(f"[DEBUG] Completed row {actual_row_num - 1}")

def add_questions_only(questions, document, add_page_break = True, images = None):
    ques_no = 1
    curr_style = 0
    if images is None:
        images = []
    
    for question in questions:
        ques, options = question["question"], question["options"]
        
        # Process question content sequentially (text, images, tables in order)
        should_page_break = add_page_break and ques_no > 1
        images, curr_style, deferred_images, deferred_difficulty = process_question_content_sequentially(ques, images, document, ques_no, should_page_break)
        
        # For questions-only document, insert deferred images after options
        for image_idx in deferred_images:
            if image_idx < len(images) and images[image_idx]:
                try:
                    print(f"  → Inserting deferred image {image_idx} after options for question {ques_no}")
                    add_image(document, images[image_idx])
                except Exception as e:
                    print(f"⚠ Warning: Could not insert deferred image {image_idx}: {e}")
        
        # For questions-only document, skip difficulty text (don't add it)
        # Difficulty text should only appear in complete documents with explanations
        if deferred_difficulty:
            print(f"  → Skipping difficulty text for questions-only document: {deferred_difficulty}")
        
        # Remove extra space after last question paragraph (before options)
        # Find the last paragraph and set space_after to 0
        if len(document.paragraphs) > 0:
            last_para = document.paragraphs[-1]
            last_para.paragraph_format.space_after = Pt(0)
        
        # Set spacing if no page break
        if not add_page_break and len(document.paragraphs) > 0:
            # Find the first paragraph we added for this question
            for para in reversed(document.paragraphs):
                if para.text.strip().startswith(f"Q{ques_no}."):
                    para.paragraph_format.space_before = Cm(1.2)
                    break

        options_idx = ["a) ", "b) ", "c) ", "d) "]
        for idx, option in enumerate(zip(options_idx, options)):
            o_paragraph = document.add_paragraph()
            # Only add space before if it's not the first option
            if idx > 0:
                o_paragraph.paragraph_format.space_before = Pt(1.1)
            else:
                o_paragraph.paragraph_format.space_before = Pt(0)  # No space before first option
            o_paragraph.paragraph_format.space_after = Pt(0)  # No space after options
            curr_style = add_text_with_style(o_paragraph, option[0] + option[1].strip(), curr_style)
                
        ques_no += 1
    return document

def set_margins(document):
    section = document.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(1.2)

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('NormalStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(font_size)
    obj_font.name = 'Source Sans Pro'

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('TableStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Arial'

    return document

def add_header_and_footer(document, folder_name):
    section = document.sections[0]
    section.footer_distance = Inches(0.1)
    section.header_distance = Inches(0.1)

    # Normalize paths for Windows
    header_path = os.path.normpath(os.path.join(folder_name, "header.png"))
    if os.path.exists(header_path):
        try:
            header_paragraph = section.header.paragraphs[0]
            header_paragraph.paragraph_format.left_indent = -Inches(0.4)
            header_paragraph.add_run().add_picture(header_path, height=Cm(2.15), width=Cm(20.88))
        except Exception:
            pass  # Silently skip if header can't be added
    
    footer_path = os.path.normpath(os.path.join(folder_name, "footer.png"))
    if os.path.exists(footer_path):
        try:
            footer_paragraph = section.footer.paragraphs[0]
            footer_paragraph.paragraph_format.left_indent = -Inches(0.5)
            footer_paragraph.add_run().add_picture(footer_path, width=Inches(8.5), height=Inches(0.48))
        except Exception:
            pass  # Silently skip if footer can't be added

def add_explanation(questions, document, images):
    answers = [-1] + list(map(lambda ques: f"{chr(ord("a") + ques["answer"])}) {ques["options"][ques["answer"]]}", questions))
    ques_no = 0
    curr_style = 0
    image_idx = 0  # Track image index
    
    for question in questions:
        ques_no += 1
        ae_paragraph = document.add_paragraph()
        ae_paragraph.paragraph_format.space_before = Cm(1.1)
        ae_paragraph.paragraph_format.line_spacing = line_spacing

        content = "\n".join([bold(f"Q{ques_no}. Answer: {answers[ques_no]}"), bold("Explanation:"), question["explanation"].rstrip("\n")])
        curr_style = add_text_with_style(ae_paragraph, content, curr_style)

        # Add images - automatically distribute from image list
        if image_idx < len(images):
            add_image(document, images[image_idx])
            image_idx += 1
        elif (question.get("image") and len(question["image"]) > 0):
            # Fallback to question-specific image
            add_image(document, question["image"])

        if question.get("source"):
            source_para = document.add_paragraph()
            add_text_with_style(source_para, question["source"], 0)


def generate_teaching_document_full(questions, folder_name, should_generate_pdfs, images, file_name):
    question_document, answer_document, question_answer_explanation_document, explanation_new_document = Document(), Document(), Document(), Document()
    explanation_only_document = Document()

    # Questions-only document should NOT have images - pass empty list
    add_questions_only(questions, set_margins(question_document), add_page_break=False, images=[])
    add_questions_and_explanation(questions, set_margins(explanation_only_document), images, page_break_before=False)
    add_answer_key(questions, set_margins(answer_document))
    # For complete document: add questions with explanations, then add answer key (no duplication)
    add_questions_and_explanation(questions, set_margins(question_answer_explanation_document), images, page_break_before=True)
    add_answer_key(questions, question_answer_explanation_document, add_page_break=True)
    add_explanation(questions, set_margins(explanation_new_document), images)

    name_suffix = {
        0: "questions_only",
        1: "answer_sheet",
        2: "complete",
        3: "explanation_and_answer_only",
        4: "explanation_only"
    }

    doc_names = []
    for i, document in enumerate([question_document, answer_document, question_answer_explanation_document, explanation_only_document, explanation_new_document]):
        add_header_and_footer(document, folder_name)
        file_name_without_ext = file_name.replace(".md", "")
        docx_path = os.path.join(folder_name, "output-docx", f"converted_{file_name_without_ext}_{name_suffix[i]}.docx")
        docx_path = os.path.normpath(docx_path)  # Normalize path for Windows
        
        # Automatically handle file saving with retries (no errors shown)
        saved = False
        for retry in range(3):
            try:
                # Auto-close Word on retry (Windows)
                if retry > 0 and is_windows():
                    close_word_if_needed()
                    time.sleep(0.5)
                
                # Remove existing file if locked
                if os.path.exists(docx_path):
                    try:
                        os.remove(docx_path)
                        time.sleep(0.2)
                    except PermissionError:
                        if retry < 2:
                            continue
                        # Last attempt: close Word and try again
                        if is_windows():
                            close_word_if_needed()
                            time.sleep(1)
                            try:
                                os.remove(docx_path)
                            except:
                                pass
                
                # Post-process: Convert any remaining markdown tables to Word tables
                try:
                    from table_converter import convert_markdown_tables_in_paragraphs
                    convert_markdown_tables_in_paragraphs(document)
                except Exception as e:
                    # If conversion fails, continue anyway
                    pass
                
                # Save the document
                document.save(docx_path)
                
                # Verify it was saved
                if os.path.exists(docx_path) and os.path.getsize(docx_path) > 0:
                    print(f"✓ Created {os.path.basename(docx_path)}")
                    saved = True
                    break
                    
            except (PermissionError, Exception):
                if retry < 2:
                    if is_windows():
                        close_word_if_needed()
                    time.sleep(1)
                    continue
                else:
                    # Silent fail - just skip this file
                    break
        
        if not saved:
            continue
        if (should_generate_pdfs):
            pdf_path = os.path.join(folder_name, "output-docx", f"converted_{file_name_without_ext}_{name_suffix[i]}.pdf")
            pdf_created = False
            
            # Try docx2pdf first
            try:
                print(f"Converting {os.path.basename(docx_path)} to PDF...")
                docx2pdf.convert(docx_path, pdf_path)
                time.sleep(0.5)
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    pdf_created = True
            except Exception as e:
                print(f"  docx2pdf failed: {e}")
            
            # Fallback: Try LibreOffice or Microsoft Word command-line tool
            if not pdf_created:
                import platform
                is_windows = platform.system() == "Windows"
                
                # Try to find LibreOffice or Microsoft Word
                soffice_paths = []
                if is_windows:
                    # Windows paths
                    program_files = os.environ.get("ProgramFiles", "C:\\Program Files")
                    program_files_x86 = os.environ.get("ProgramFiles(x86)", "C:\\Program Files (x86)")
                    soffice_paths = [
                        os.path.join(program_files, "LibreOffice", "program", "soffice.exe"),
                        os.path.join(program_files_x86, "LibreOffice", "program", "soffice.exe"),
                        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
                        "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
                        "soffice.exe",  # In PATH
                        "soffice"  # In PATH
                    ]
                    # Also try Microsoft Word (if available)
                    word_paths = [
                        os.path.join(program_files, "Microsoft Office", "root", "Office16", "WINWORD.EXE"),
                        os.path.join(program_files_x86, "Microsoft Office", "root", "Office16", "WINWORD.EXE"),
                        "C:\\Program Files\\Microsoft Office\\root\\Office16\\WINWORD.EXE",
                        "C:\\Program Files (x86)\\Microsoft Office\\root\\Office16\\WINWORD.EXE",
                    ]
                else:
                    # macOS/Linux paths
                    soffice_paths = [
                        "/opt/homebrew/bin/soffice",
                        "/usr/local/bin/soffice",
                        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                        "soffice"  # In PATH
                    ]
                    word_paths = []
                
                soffice = None
                word_exe = None
                
                # Try to find LibreOffice
                for path in soffice_paths:
                    if os.path.exists(path) or (path.endswith("soffice") or path.endswith("soffice.exe")):
                        try:
                            if is_windows:
                                result = subprocess.run([path, "--version"], capture_output=True, timeout=5, shell=True)
                            else:
                                result = subprocess.run([path, "--version"], capture_output=True, timeout=2)
                            if result.returncode == 0:
                                soffice = path
                                break
                        except:
                            continue
                
                # Try to find Microsoft Word (Windows only)
                if is_windows and not soffice:
                    for path in word_paths:
                        if os.path.exists(path):
                            word_exe = path
                            break
                
                if soffice:
                    try:
                        print(f"  Trying LibreOffice conversion...")
                        # Use LibreOffice headless mode to convert
                        # Normalize paths for Windows
                        docx_path_normalized = os.path.normpath(docx_path)
                        pdf_dir_normalized = os.path.normpath(os.path.dirname(pdf_path))
                        
                        if is_windows:
                            # Windows: Use shell=True and proper path handling
                            cmd = [
                                soffice,
                                "--headless",
                                "--convert-to", "pdf",
                                "--outdir", pdf_dir_normalized,
                                docx_path_normalized
                            ]
                            result = subprocess.run(cmd, timeout=60, capture_output=True, shell=True, text=True)
                        else:
                            # macOS/Linux
                            result = subprocess.run([
                                soffice,
                                "--headless",
                                "--convert-to", "pdf",
                                "--outdir", pdf_dir_normalized,
                                docx_path_normalized
                            ], timeout=30, capture_output=True, text=True)
                        
                        if result.returncode != 0:
                            print(f"  LibreOffice returned error code: {result.returncode}")
                            if result.stderr:
                                error_msg = result.stderr[:300] if result.stderr else "Unknown error"
                                print(f"  Error: {error_msg}")
                        
                        # LibreOffice creates PDF with same name but .pdf extension
                        time.sleep(1)  # Give it more time on Windows
                        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                            pdf_created = True
                        else:
                            # Sometimes LibreOffice uses a different naming convention
                            base_name = os.path.splitext(os.path.basename(docx_path))[0]
                            alt_pdf_path = os.path.join(os.path.dirname(pdf_path), f"{base_name}.pdf")
                            if os.path.exists(alt_pdf_path) and os.path.getsize(alt_pdf_path) > 0:
                                os.rename(alt_pdf_path, pdf_path)
                                pdf_created = True
                    except subprocess.TimeoutExpired:
                        print(f"  LibreOffice conversion timed out (file may be too large)")
                    except Exception as e:
                        print(f"  LibreOffice conversion failed: {e}")
                        print(f"  Make sure LibreOffice is properly installed and 'soffice' is in your PATH")
            
            if pdf_created:
                doc_names.append(pdf_path)
                print(f"✓ Created PDF: {os.path.basename(pdf_path)}")
            else:
                # Silent fail - DOCX files are already created, PDF is optional
                print(f"⚠ PDF conversion skipped for {os.path.basename(docx_path)}")
                print(f"  (DOCX file is available - you can convert manually if needed)")

    if (should_generate_pdfs):
        convert_all_pdfs(doc_names, folder_name)

should_generate_pdfs = input(">> Generate final pdfs ?\n")
if (should_generate_pdfs.lower().strip().startswith("y")):
    should_generate_pdfs = True
else: 
    should_generate_pdfs = False

folder_input = input(">> Enter folder which contains header.png, footer.png, first_and_last_page.pdf (or '.' for current directory)\n").strip()
if folder_input == "." or folder_input == "":
    folder_name = os.path.dirname(os.path.abspath(__file__)) or os.getcwd()
else:
    # Expand user path if needed
    folder_input = os.path.expanduser(folder_input)
    # Convert to absolute path
    folder_name = os.path.abspath(folder_input)

font_size = int(input(">> Enter font-size\n"))
line_spacing = float(input(">> Enter line-spacing\n"))

# Normalize path for Windows
folder_name = os.path.normpath(folder_name)

# Auto-close Word before creating directories (Windows)
if is_windows():
    close_word_if_needed()

# Create output directories with automatic retry
output_dir = os.path.join(folder_name, "output")
output_docx_dir = os.path.join(folder_name, "output-docx")

for retry in range(3):
    try:
        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(output_docx_dir, exist_ok=True)
        break
    except PermissionError:
        if retry < 2 and is_windows():
            close_word_if_needed()
            time.sleep(1)
        else:
            print(f"Error: Cannot create output folders. Please check permissions.")
            exit(1)
    except Exception as e:
        if retry < 2:
            time.sleep(0.5)
        else:
            print(f"Error creating output folders: {e}")
            exit(1)

files_to_convert_dir = os.path.join(folder_name, "files_to_convert")
if not os.path.exists(files_to_convert_dir):
    print(f"Error: Directory '{files_to_convert_dir}' does not exist.")
    print(f"Please create it and add your .md, .doc, or .docx files there.")
    exit(1)

# Supported file extensions
supported_extensions = ['.md', '.doc', '.docx']

for file_name in os.listdir(files_to_convert_dir):
    file_ext = os.path.splitext(file_name)[1].lower()
    if file_ext not in supported_extensions:
        continue

    file_path = os.path.join(files_to_convert_dir, file_name)
    print(f"\nProcessing: {file_name}")

    try:
        # Read file content (handles MD, DOC, DOCX) - returns (lines, images) tuple
        lines, file_images = read_file_content(file_path)
        
        # Process the content
        parsing_state = ParsingState()
        # Join lines and process similar to markdown
        content = "".join(lines).replace("***", "**")
        lines = content.splitlines(keepends=True)
        lines = list(filter(lambda line: not (len(set(line)) == 1 and line[0] == "\n"), lines))
        lines = list(map(lambda line: line.replace("\\", ""), lines))
        
        # Process lines, but tables will be extracted separately by parse_question_and_table
        for line in lines:
            parsing_state.set_or_update_state(line)

        parsing_state.flush_state()      
        questions = parsing_state.questions
        
        # Merge images from file with images from parsing
        # File images (from DOCX) take priority
        if file_images:
            images = file_images
        else:
            images = parsing_state.images
        
        if len(questions) == 0:
            print(f"⚠ Warning: No questions found in {file_name}")
            print(f"  Make sure the file follows the expected format:")
            print(f"  Question: ...")
            print(f"  a) ...")
            print(f"  b) ...")
            print(f"  c) ...")
            print(f"  d) ...")
            print(f"  Answer: ...")
            print(f"  Explanation: ...")
            continue
        
        # Filter out questions without answers before processing
        valid_questions = [q for q in questions if q["answer"] != -1]
        invalid_count = len(questions) - len(valid_questions)
        
        if invalid_count > 0:
            print(f"⚠ Warning: {invalid_count} question(s) skipped due to missing answers")
            print(f"  Tip: Make sure answers are in format: 'Answer: a' or 'Answer: [option text]'")
        
        if len(valid_questions) == 0:
            print(f"✗ No valid questions found in {file_name}")
            continue
        
        # Set default explanation for questions without it
        for q in valid_questions:
            if not q.get("explanation") or len(q["explanation"].strip()) == 0:
                q["explanation"] = "No explanation provided."
        
        check_state(valid_questions)
        
        # Use base filename without extension for output
        base_file_name = os.path.splitext(file_name)[0] + ".md"
        generate_classplus_table_formatted_document(base_file_name, valid_questions, folder_name)
        generate_teaching_document_full(valid_questions, folder_name, should_generate_pdfs, images, base_file_name)
        
        print(f"✓ Successfully processed {file_name}")
        
    except Exception as e:
        print(f"✗ Error processing {file_name}: {e}")
        print(f"  Skipping this file...")
        continue