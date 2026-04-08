#!/usr/bin/env python3
"""
Script to create demo DOC and DOCX files with sample questions
"""
from docx import Document
import os

# Sample questions content
questions_content = """Question: What is the capital of France?
a) London
b) Berlin
c) Paris
d) Madrid
Answer: c
Explanation: Paris is the capital and largest city of France. It is located in the north-central part of the country on the Seine River.
Source: Geography Basics - Chapter 1

Question: Which of the following is the largest planet in our solar system?
a) Earth
b) Jupiter
c) Saturn
d) Neptune
Answer: b
Explanation: Jupiter is the largest planet in our solar system. It is a gas giant with a mass greater than all other planets combined.
Source: Astronomy Fundamentals - Chapter 3

Question: What is 2 + 2?
| Operation | Result |
|----------|--------|
| Addition | 4 |
| Subtraction | 0 |
a) 3
b) 4
c) 5
d) 6
Answer: b
Explanation: The sum of 2 and 2 equals 4. This is a basic arithmetic operation.
Source: Mathematics Basics
"""

def create_docx_file():
    """Create a DOCX file with sample questions"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sample Questions', 0)
    
    # Split content into lines and add to document
    lines = questions_content.split('\n')
    for line in lines:
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()  # Empty line
    
    # Save DOCX file
    output_dir = "files_to_convert"
    os.makedirs(output_dir, exist_ok=True)
    docx_path = os.path.join(output_dir, "demo_questions.docx")
    doc.save(docx_path)
    print(f"✓ Created: {docx_path}")
    return docx_path

def create_doc_file():
    """Create a DOC file by converting DOCX using LibreOffice"""
    import subprocess
    import platform
    
    docx_path = os.path.join("files_to_convert", "demo_questions.docx")
    doc_path = os.path.join("files_to_convert", "demo_questions.doc")
    
    if not os.path.exists(docx_path):
        print("✗ DOCX file not found. Creating it first...")
        create_docx_file()
    
    is_windows = platform.system() == "Windows"
    
    # Find LibreOffice
    if is_windows:
        program_files = os.environ.get("ProgramFiles", "C:\\Program Files")
        soffice = os.path.join(program_files, "LibreOffice", "program", "soffice.exe")
        if not os.path.exists(soffice):
            soffice = "soffice.exe"
    else:
        soffice = "soffice"
    
    try:
        # Convert DOCX to DOC
        if is_windows:
            result = subprocess.run([
                soffice,
                "--headless",
                "--convert-to", "doc",
                "--outdir", "files_to_convert",
                docx_path
            ], capture_output=True, timeout=30, shell=True)
        else:
            result = subprocess.run([
                soffice,
                "--headless",
                "--convert-to", "doc",
                "--outdir", "files_to_convert",
                docx_path
            ], capture_output=True, timeout=30)
        
        if result.returncode == 0 and os.path.exists(doc_path):
            print(f"✓ Created: {doc_path}")
            return doc_path
        else:
            print(f"⚠ Could not create DOC file automatically.")
            print(f"  LibreOffice conversion failed or not installed.")
            print(f"  You can manually convert {docx_path} to DOC format if needed.")
            return None
    except Exception as e:
        print(f"⚠ Could not create DOC file: {e}")
        print(f"  DOCX file is available: {docx_path}")
        print(f"  Note: DOC format requires LibreOffice for conversion.")
        return None

if __name__ == "__main__":
    print("Creating demo files...")
    print()
    
    # Create DOCX file
    docx_path = create_docx_file()
    
    # Try to create DOC file
    print()
    print("Attempting to create DOC file (requires LibreOffice)...")
    doc_path = create_doc_file()
    
    print()
    print("=" * 60)
    print("Demo files created!")
    print("=" * 60)
    print(f"✓ DOCX file: {docx_path}")
    if doc_path:
        print(f"✓ DOC file: {doc_path}")
    else:
        print(f"⚠ DOC file: Not created (LibreOffice required)")
    print()
    print("You can now run: python docx_generator.py")
    print("=" * 60)






