"""
Post-processing utility to convert markdown table text into actual Word tables
"""
import re
from docx.shared import Pt, Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_text_with_style(paragraph, text: str, curr_style):
    """Helper to add text with style - simplified version"""
    # Parse markdown-style formatting
    texts = []
    curr_text = ""
    i = 0
    
    while i < len(text):
        ch = text[i]
        if ch == "*" and i + 1 < len(text) and text[i + 1] == "*":
            if len(curr_text) > 0:
                texts.append({"text": curr_text, "bold": False})
                curr_text = ""
            # Toggle bold
            i += 2
            while i < len(text) and text[i] == "*":
                i += 1
            # Find closing **
            while i < len(text):
                if text[i] == "*" and i + 1 < len(text) and text[i + 1] == "*":
                    texts.append({"text": curr_text, "bold": True})
                    curr_text = ""
                    i += 2
                    break
                else:
                    curr_text += text[i]
                    i += 1
            continue
        else:
            curr_text += ch
        i += 1
    
    if len(curr_text) > 0:
        texts.append({"text": curr_text, "bold": False})
    
    for text_obj in texts:
        run = paragraph.add_run(text_obj["text"])
        run.bold = text_obj["bold"]
        # Explicitly ensure bold is set in XML for better PDF compatibility
        if text_obj["bold"]:
            try:
                if hasattr(run, '_element'):
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        run._element.insert(0, rPr)
                    b_elem = rPr.find(qn('w:b'))
                    if b_elem is None:
                        b_elem = OxmlElement('w:b')
                        rPr.append(b_elem)
                    b_elem.set(qn('w:val'), 'true')
            except:
                pass  # If XML manipulation fails, run.bold = True should still work
    
    return curr_style

def detect_and_convert_markdown_tables(document):
    """
    Scan document for markdown table patterns and convert them to actual Word tables.
    Pattern: Lines starting with | that form a table structure
    """
    # Collect all paragraphs
    paragraphs_to_check = []
    for para in document.paragraphs:
        paragraphs_to_check.append(para)
    
    # Find markdown table patterns
    i = 0
    while i < len(paragraphs_to_check):
        para = paragraphs_to_check[i]
        text = para.text.strip()
        
        # Check if this looks like a table row
        if text.startswith("|") and "|" in text[1:]:
            # Found potential table start - collect consecutive table rows
            table_rows = []
            start_idx = i
            current_idx = i
            
            # Collect all consecutive table-like lines
            while current_idx < len(paragraphs_to_check):
                current_para = paragraphs_to_check[current_idx]
                current_text = current_para.text.strip()
                
                # Check if it's a table row or separator
                if current_text.startswith("|"):
                    # Check if it's a separator (all dashes/colons)
                    is_separator = all(c in "-: |" for c in current_text.replace(" ", ""))
                    if not is_separator:
                        # Parse table row - preserve empty cells to maintain column alignment
                        # Split by | and take elements from index 1 to -1 (excluding first/last empty strings)
                        parts = current_text.split("|")
                        cells = [part.strip() for part in parts[1:-1]]  # Preserve all cells including empty ones
                        if len(cells) > 0:
                            table_rows.append(cells)
                    # Move to next paragraph
                    current_idx += 1
                else:
                    # Not a table row, stop collecting
                    break
            
            # If we found table rows, convert them
            if len(table_rows) >= 2:  # At least header + 1 data row
                # Remove the markdown paragraphs
                for j in range(start_idx, current_idx):
                    if j < len(paragraphs_to_check):
                        # Clear the paragraph text
                        paragraphs_to_check[j].clear()
                
                # Insert table before the first removed paragraph
                if start_idx < len(paragraphs_to_check):
                    # Create table
                    num_cols = max(len(row) for row in table_rows) if table_rows else 0
                    if num_cols > 0:
                        # Find the parent element to insert table
                        parent = paragraphs_to_check[start_idx]._element.getparent()
                        
                        # Create table element
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls, qn
                        
                        # Create table XML
                        tbl_xml = f'<w:tbl {nsdecls("w")}><w:tblPr><w:tblStyle w:val="TableGrid"/><w:tblW w:w="0" w:type="auto"/></w:tblPr><w:tblGrid>'
                        for _ in range(num_cols):
                            tbl_xml += '<w:gridCol w:w="2000"/>'
                        tbl_xml += '</w:tblGrid>'
                        
                        # Add rows
                        for row in table_rows:
                            tbl_xml += '<w:tr>'
                            for cell_idx in range(num_cols):
                                cell_text = row[cell_idx] if cell_idx < len(row) else ""
                                tbl_xml += f'<w:tc><w:tcPr><w:tcW w:w="2000" w:type="dxa"/></w:tcPr><w:p><w:r><w:t>{cell_text}</w:t></w:r></w:p></w:tc>'
                            tbl_xml += '</w:tr>'
                        
                        tbl_xml += '</w:tbl>'
                        
                        # Parse and insert
                        tbl_element = parse_xml(tbl_xml)
                        parent.insert_before(tbl_element, paragraphs_to_check[start_idx]._element)
                
                # Skip the paragraphs we just processed
                i = current_idx
                continue
        
        i += 1

def convert_markdown_tables_in_paragraphs(document):
    """
    Find paragraphs with table markdown and convert them to actual Word tables.
    This is a post-processing step that runs after document creation.
    """
    # Get all paragraphs as a list (we need to work with indices)
    paras = list(document.paragraphs)
    
    i = 0
    while i < len(paras):
        para = paras[i]
        text = para.text.strip()
        
        # Check if this paragraph contains table markdown
        if text.startswith("|") and text.count("|") >= 2:
            # Collect table rows starting from this paragraph
            table_rows = []
            rows_to_clear = []
            j = i
            
            # Collect consecutive table rows
            while j < len(paras):
                row_text = paras[j].text.strip()
                if row_text.startswith("|"):
                    # Check if separator row (all dashes/colons)
                    is_separator = all(c in "-: |" for c in row_text.replace(" ", ""))
                    if not is_separator:
                        # Parse cells - preserve empty cells to maintain column alignment
                        # Split by | and take elements from index 1 to -1 (excluding first/last empty strings)
                        parts = row_text.split("|")
                        cells = [part.strip() for part in parts[1:-1]]  # Preserve all cells including empty ones
                        if len(cells) > 0:
                            table_rows.append(cells)
                            rows_to_clear.append(j)
                    j += 1
                else:
                    # Not a table row, stop collecting
                    break
            
            # If we have a valid table (at least 2 rows), create it
            if len(table_rows) >= 2:
                # Determine column count
                num_cols = max(len(row) for row in table_rows) if table_rows else 0
                
                if num_cols > 0:
                    # Create table before the first markdown paragraph
                    table = document.add_table(rows=0, cols=num_cols)
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    table.autofit = False
                    
                    # Set column widths
                    col_width = Inches(1.8) / num_cols
                    
                    # Add rows to table
                    for row_idx, row in enumerate(table_rows):
                        table_row = table.add_row()
                        table_row.height = Cm(1.1)
                        cells = table_row.cells
                        
                        # Set column width for first row
                        if row_idx == 0:
                            for cell in cells:
                                cell.width = col_width
                        
                        for idx in range(num_cols):
                            cell_text = row[idx] if idx < len(row) else ""
                            cell_para = cells[idx].paragraphs[0]
                            cell_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            cell_para.paragraph_format.space_before = Pt(3)
                            cell_para.paragraph_format.space_after = Pt(3)
                            
                            if cell_text:
                                add_text_with_style(cell_para, cell_text, 0)
                            else:
                                cell_para.text = ""
                    
                    # Clear the markdown paragraphs (remove the text)
                    for idx in rows_to_clear:
                        if idx < len(paras):
                            # Clear all runs in the paragraph
                            for run in paras[idx].runs:
                                run.text = ""
                            paras[idx].text = ""
                    
                    # Move table element to correct position (before first cleared para)
                    # Get table element
                    table_element = table._element
                    # Get parent
                    parent = table_element.getparent()
                    # Get first cleared paragraph element
                    if rows_to_clear and rows_to_clear[0] < len(paras):
                        first_para_element = paras[rows_to_clear[0]]._element
                        # Move table before first paragraph
                        parent.insert_before(table_element, first_para_element)
                
                # Skip processed paragraphs
                i = j
                continue
        
        i += 1

