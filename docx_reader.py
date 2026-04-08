"""
Functions to read and extract text from DOC and DOCX files with formatting preservation
"""
import os
import shutil
import base64
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def is_run_bold(run):
    """
    Comprehensive check for bold formatting in a run.
    Checks multiple sources:
    1. Direct run.bold property (True/False/None)
    2. Run's character style if run.bold is None
    3. Paragraph's style if still None
    4. XML element directly for <w:b> tag
    """
    # Method 1: Check direct property
    if run.bold is True:
        return True
    if run.bold is False:
        return False
    
    # Method 2: Check run's character style
    if hasattr(run, 'style') and run.style:
        try:
            if hasattr(run.style, 'font') and run.style.font:
                if run.style.font.bold is True:
                    return True
                if run.style.font.bold is False:
                    return False
        except:
            pass
    
    # Method 3: Check paragraph's style
    if hasattr(run, 'parent') and hasattr(run.parent, 'style'):
        try:
            para_style = run.parent.style
            if para_style and hasattr(para_style, 'font') and para_style.font:
                if para_style.font.bold is True:
                    return True
                if para_style.font.bold is False:
                    return False
        except:
            pass
    
    # Method 4: Check XML directly for <w:b> element
    try:
        if hasattr(run, '_element'):
            element = run._element
            # Look for <w:rPr><w:b/> or <w:rPr><w:b w:val="true"/>
            rPr = element.find(qn('w:rPr'))
            if rPr is not None:
                b_elem = rPr.find(qn('w:b'))
                if b_elem is not None:
                    # Check if bold is explicitly set
                    val = b_elem.get(qn('w:val'))
                    if val is None or val.lower() in ('true', '1', 'on'):
                        return True
                    if val.lower() in ('false', '0', 'off'):
                        return False
    except:
        pass
    
    # Default: not bold if we can't determine
    return False

def is_run_italic(run):
    """
    Comprehensive check for italic formatting in a run.
    Similar to is_run_bold but for italic.
    """
    # Method 1: Check direct property
    if run.italic is True:
        return True
    if run.italic is False:
        return False
    
    # Method 2: Check run's character style
    if hasattr(run, 'style') and run.style:
        try:
            if hasattr(run.style, 'font') and run.style.font:
                if run.style.font.italic is True:
                    return True
                if run.style.font.italic is False:
                    return False
        except:
            pass
    
    # Method 3: Check paragraph's style
    if hasattr(run, 'parent') and hasattr(run.parent, 'style'):
        try:
            para_style = run.parent.style
            if para_style and hasattr(para_style, 'font') and para_style.font:
                if para_style.font.italic is True:
                    return True
                if para_style.font.italic is False:
                    return False
        except:
            pass
    
    # Method 4: Check XML directly for <w:i> element
    try:
        if hasattr(run, '_element'):
            element = run._element
            rPr = element.find(qn('w:rPr'))
            if rPr is not None:
                i_elem = rPr.find(qn('w:i'))
                if i_elem is not None:
                    val = i_elem.get(qn('w:val'))
                    if val is None or val.lower() in ('true', '1', 'on'):
                        return True
                    if val.lower() in ('false', '0', 'off'):
                        return False
    except:
        pass
    
    # Default: not italic if we can't determine
    return False

def extract_formatted_text_from_runs(runs):
    """
    Extract text from runs with formatting (bold, italic) converted to markdown style
    Uses comprehensive detection that checks run properties, styles, and XML directly
    """
    formatted_text = ""
    for run in runs:
        text = run.text
        if not text:
            continue
        
        # Use comprehensive detection functions
        is_bold = is_run_bold(run)
        is_italic = is_run_italic(run)
        
        if is_bold and is_italic:
            formatted_text += f"***{text}***"
        elif is_bold:
            formatted_text += f"**{text}**"
        elif is_italic:
            formatted_text += f"*{text}*"
        else:
            formatted_text += text
    
    return formatted_text

def extract_table_as_markdown(table):
    """
    Extract table and convert to markdown table format
    """
    lines = []
    for row_idx, row in enumerate(table.rows):
        row_cells = []
        for cell in row.cells:
            # Extract formatted text from cell
            cell_text = ""
            for para_idx, paragraph in enumerate(cell.paragraphs):
                if paragraph.runs:
                    cell_text += extract_formatted_text_from_runs(paragraph.runs)
                else:
                    cell_text += paragraph.text
                # Add space between multiple paragraphs in same cell
                if para_idx < len(cell.paragraphs) - 1:
                    cell_text += " "
            row_cells.append(cell_text.strip())
        
        if any(row_cells):  # Only add non-empty rows
            lines.append("| " + " | ".join(row_cells) + " |")
            
            # Add header separator after first row (markdown table format)
            if row_idx == 0:
                num_cols = len(row_cells)
                separator = "| " + " | ".join(["---"] * num_cols) + " |"
                lines.append(separator)
    
    return lines

def extract_images_from_docx(docx_path):
    """
    Extract images from DOCX file and return as list of base64 encoded images
    """
    images = []
    try:
        from zipfile import ZipFile
        
        # DOCX files are ZIP archives
        with ZipFile(docx_path, 'r') as docx_zip:
            # Find image files in word/media directory
            image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/') and 
                          f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
            
            for img_file in image_files:
                try:
                    img_data = docx_zip.read(img_file)
                    # Determine image type
                    img_ext = os.path.splitext(img_file)[1].lower()
                    if img_ext == '.png':
                        mime_type = 'image/png'
                    elif img_ext in ['.jpg', '.jpeg']:
                        mime_type = 'image/jpeg'
                    elif img_ext == '.gif':
                        mime_type = 'image/gif'
                    else:
                        mime_type = 'image/png'
                    
                    # Convert to base64
                    base64_str = base64.b64encode(img_data).decode('utf-8')
                    images.append(f"data:{mime_type};base64,{base64_str}")
                except Exception as e:
                    # Skip images that can't be read
                    continue
    except Exception as e:
        # If image extraction fails, continue without images
        pass
    
    return images

def extract_images_from_paragraph(paragraph):
    """
    Extract images from a paragraph and return as list of base64 encoded images
    Uses multiple methods to find images in runs
    """
    images = []
    try:
        from docx.oxml.ns import qn
        
        for run in paragraph.runs:
            if not hasattr(run, '_element'):
                continue
                
            element = run._element
            
            # Method 1: Find all blip elements using various namespaces
            blip_elements = []
            
            # Try different namespace patterns
            namespaces = [
                ('a', 'http://schemas.openxmlformats.org/drawingml/2006/main'),
                ('pic', 'http://schemas.openxmlformats.org/drawingml/2006/picture'),
                ('r', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships')
            ]
            
            # Search for blip elements with different approaches
            try:
                # Direct xpath with namespace
                blip_elements = element.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
            except:
                pass
            
            if not blip_elements:
                # Try findall with full namespace
                blip_elements = element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            
            if not blip_elements:
                # Try with qn
                try:
                    blip_elements = element.findall(qn('a:blip'))
                except:
                    pass
            
            # Process each blip element found
            for blip in blip_elements:
                try:
                    # Get relationship ID
                    rId = blip.get(qn('r:embed'))
                    if not rId:
                        # Try alternative attribute name
                        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    
                    if rId and hasattr(run, 'part') and rId in run.part.related_parts:
                        # Get the image part
                        image_part = run.part.related_parts[rId]
                        image_data = image_part.blob
                        
                        # Determine image type
                        content_type = image_part.content_type
                        if 'png' in content_type:
                            mime_type = 'image/png'
                        elif 'jpeg' in content_type or 'jpg' in content_type:
                            mime_type = 'image/jpeg'
                        elif 'gif' in content_type:
                            mime_type = 'image/gif'
                        else:
                            mime_type = 'image/png'
                        
                        # Convert to base64
                        base64_str = base64.b64encode(image_data).decode('utf-8')
                        images.append(f"data:{mime_type};base64,{base64_str}")
                except Exception as e:
                    continue
            
            # Method 2: Check for inline shapes (alternative approach)
            if not blip_elements:
                try:
                    # Look for drawing elements
                    drawings = element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                    for drawing in drawings:
                        # Try to find blip in drawing
                        blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                        for blip in blips:
                            try:
                                rId = blip.get(qn('r:embed'))
                                if rId and hasattr(run, 'part') and rId in run.part.related_parts:
                                    image_part = run.part.related_parts[rId]
                                    image_data = image_part.blob
                                    content_type = image_part.content_type
                                    mime_type = 'image/png'
                                    if 'png' in content_type:
                                        mime_type = 'image/png'
                                    elif 'jpeg' in content_type or 'jpg' in content_type:
                                        mime_type = 'image/jpeg'
                                    elif 'gif' in content_type:
                                        mime_type = 'image/gif'
                                    base64_str = base64.b64encode(image_data).decode('utf-8')
                                    images.append(f"data:{mime_type};base64,{base64_str}")
                            except:
                                continue
                except:
                    pass
                    
    except Exception as e:
        pass
    return images

def extract_text_from_docx(docx_path):
    """
    Extract text from a DOCX file with formatting, images, and tables preserved
    Returns: (lines, images) tuple where images is a list of (image_data, position) tuples
    Processes elements in document order to preserve question order
    """
    try:
        doc = Document(docx_path)
        lines = []
        images = []  # Will store (image_data, line_index) tuples for position tracking
        
        # Track which paragraphs are inside tables to avoid double processing
        table_paragraphs = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        table_paragraphs.add(id(paragraph))
        
        # Create mappings for quick lookup
        element_to_para = {para._element: para for para in doc.paragraphs}
        element_to_table = {table._element: table for table in doc.tables}
        
        # Track which tables we've processed
        processed_tables = set()
        
        # Process elements in document order by iterating through body XML elements
        # This preserves the original order of paragraphs and tables
        try:
            # Get the body element which contains all paragraphs and tables in order
            body = doc._body._body
            
            # Iterate through child elements in document order
            for element in body:
                element_tag = element.tag
                
                # Check if it's a paragraph (tag ends with '}p')
                if element_tag.endswith('}p'):
                    para = element_to_para.get(element)
                    if para and id(para) not in table_paragraphs:
                        # Extract images from this paragraph first (before text)
                        para_images = extract_images_from_paragraph(para)
                        for img in para_images:
                            # Insert image marker in the lines at current position
                            lines.append(f"[IMAGE:{len(images)}]\n")
                            images.append(img)
                        
                        # Extract formatted text from runs
                        if para.runs:
                            formatted_text = extract_formatted_text_from_runs(para.runs)
                        else:
                            formatted_text = para.text
                        
                        if formatted_text.strip():
                            lines.append(formatted_text + "\n")
                
                # Check if it's a table (tag ends with '}tbl')
                elif element_tag.endswith('}tbl'):
                    table = element_to_table.get(element)
                    if table and id(table) not in processed_tables:
                        processed_tables.add(id(table))
                        table_lines = extract_table_as_markdown(table)
                        if table_lines:
                            lines.extend(table_lines)
                            lines.append("\n")  # Add blank line after table
        except (AttributeError, KeyError, TypeError) as e:
            # If ordered extraction fails, fall back to simple method
            for paragraph in doc.paragraphs:
                if id(paragraph) not in table_paragraphs:
                    # Extract images from paragraph
                    para_images = extract_images_from_paragraph(paragraph)
                    for img in para_images:
                        lines.append(f"[IMAGE:{len(images)}]\n")
                        images.append(img)
                    
                    if paragraph.runs:
                        formatted_text = extract_formatted_text_from_runs(paragraph.runs)
                    else:
                        formatted_text = paragraph.text
                    
                    if formatted_text.strip():
                        lines.append(formatted_text + "\n")
            
            for table in doc.tables:
                table_lines = extract_table_as_markdown(table)
                if table_lines:
                    lines.extend(table_lines)
                    lines.append("\n")
        
        # Fallback: if no images found with position tracking, use simple extraction
        if not images:
            images = extract_images_from_docx(docx_path)
        
        return lines, images
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {e}")

def extract_text_from_doc(doc_path):
    """
    Extract text from a DOC file (older format)
    Note: This requires LibreOffice or conversion
    """
    # Try to convert DOC to DOCX first using LibreOffice
    import subprocess
    import tempfile
    import platform
    
    is_windows = platform.system() == "Windows"
    
    # Find LibreOffice
    if is_windows:
        program_files = os.environ.get("ProgramFiles", "C:\\Program Files")
        soffice = os.path.join(program_files, "LibreOffice", "program", "soffice.exe")
        if not os.path.exists(soffice):
            soffice = "soffice.exe"
    else:
        soffice = "soffice"
    
    try:
        # Convert DOC to DOCX using LibreOffice
        temp_dir = tempfile.mkdtemp()
        temp_docx = os.path.join(temp_dir, "temp.docx")
        
        if is_windows:
            result = subprocess.run([
                soffice,
                "--headless",
                "--convert-to", "docx",
                "--outdir", temp_dir,
                doc_path
            ], capture_output=True, timeout=30, shell=True)
        else:
            result = subprocess.run([
                soffice,
                "--headless",
                "--convert-to", "docx",
                "--outdir", temp_dir,
                doc_path
            ], capture_output=True, timeout=30)
        
        # Find the converted file (LibreOffice might use different naming)
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        # Try different possible names
        possible_names = [
            os.path.join(temp_dir, f"{base_name}.docx"),
            os.path.join(temp_dir, "temp.docx"),
            os.path.join(temp_dir, os.path.basename(doc_path).replace('.doc', '.docx'))
        ]
        
        converted_file = None
        for name in possible_names:
            if os.path.exists(name):
                converted_file = name
                break
        
        if converted_file and os.path.exists(converted_file):
            lines, images = extract_text_from_docx(converted_file)
            # Clean up
            try:
                os.remove(converted_file)
                os.rmdir(temp_dir)
            except:
                pass
            return lines, images
        else:
            # Clean up temp directory
            try:
                import shutil
                shutil.rmtree(temp_dir)
            except:
                pass
            raise Exception("Could not convert DOC to DOCX. Please install LibreOffice or convert the file manually.")
            
    except Exception as e:
        # Fallback: try to read as text (won't work well but better than nothing)
        try:
            with open(doc_path, 'rb') as f:
                content = f.read()
                # Try to extract readable text (basic approach)
                text = content.decode('utf-8', errors='ignore')
                lines = [line + "\n" for line in text.split('\n') if line.strip()]
                return lines, []
        except:
            raise Exception(f"Could not read DOC file. Please convert it to DOCX first. Error: {e}")

def read_file_content(file_path):
    """
    Read content from a file (MD, DOCX, or DOC) and return as list of lines
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.md':
        # Read markdown file
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            return content.splitlines(keepends=True)
    elif file_ext == '.docx':
        # Read DOCX file
        return extract_text_from_docx(file_path)
    elif file_ext == '.doc':
        # Read DOC file (older format)
        return extract_text_from_doc(doc_path)
    else:
        raise Exception(f"Unsupported file format: {file_ext}")

